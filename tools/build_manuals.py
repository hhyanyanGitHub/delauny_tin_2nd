"""Generate the two dterrain Word manuals.

Run with the bundled workspace Python that provides python-docx.
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "manuals"
VERSION = "0.27.0"
TODAY = date(2026, 7, 18).isoformat()

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "203748"
MUTED = "66727D"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
INK = "20262C"
GREEN = "2F6B4F"
GOLD = "7A5A00"
RED = "9B1C1C"


def set_run_font(run, size=None, bold=None, italic=None, color=None,
                 latin="Calibri", east_asia="Microsoft YaHei"):
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start),
                       ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{tag}"))
        if element is None:
            element = OxmlElement(f"w:{tag}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def keep_table_row_together(row):
    """Prevent Word from splitting one logical table row across pages."""
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


class Manual:
    def __init__(self, title, short_title, subtitle, audience):
        self.doc = Document()
        self.title = title
        self.short_title = short_title
        self.subtitle = subtitle
        self.audience = audience
        self._number_counter = 0
        self._configure_document()

    def _configure_document(self):
        doc = self.doc
        section = doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        section.different_first_page_header_footer = True

        styles = doc.styles
        normal = styles["Normal"]
        normal.font.name = "Calibri"
        normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        normal.font.size = Pt(11)
        normal.font.color.rgb = RGBColor.from_string(INK)
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.25
        normal.paragraph_format.widow_control = True

        for name, size, color, before, after in (
            ("Heading 1", 16, BLUE, 18, 10),
            ("Heading 2", 13, BLUE, 14, 7),
            ("Heading 3", 12, DARK_BLUE, 10, 5),
        ):
            style = styles[name]
            style.font.name = "Calibri"
            style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
            style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            style.font.size = Pt(size)
            style.font.bold = True
            style.font.color.rgb = RGBColor.from_string(color)
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)
            style.paragraph_format.keep_with_next = True
            style.paragraph_format.widow_control = True

        for name in ("List Bullet", "List Number"):
            style = styles[name]
            style.font.name = "Calibri"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            style.font.size = Pt(11)
            pf = style.paragraph_format
            pf.left_indent = Inches(0.375)
            pf.first_line_indent = Inches(-0.188)
            pf.space_after = Pt(4)
            pf.line_spacing = 1.25
            pf.widow_control = True

        code = styles.add_style("Dterrain Code", WD_STYLE_TYPE.PARAGRAPH)
        code.font.name = "Consolas"
        code._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
        code._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
        code._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        code.font.size = Pt(8.5)
        code.font.color.rgb = RGBColor.from_string("1E2933")
        code.paragraph_format.left_indent = Inches(0.18)
        code.paragraph_format.right_indent = Inches(0.18)
        code.paragraph_format.space_before = Pt(3)
        code.paragraph_format.space_after = Pt(7)
        code.paragraph_format.line_spacing = 1.0
        code.paragraph_format.keep_together = True

        caption = styles["Caption"]
        caption.font.name = "Calibri"
        caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        caption.font.size = Pt(9)
        caption.font.italic = True
        caption.font.color.rgb = RGBColor.from_string(MUTED)
        caption.paragraph_format.space_before = Pt(4)
        caption.paragraph_format.space_after = Pt(4)

        self._set_running_furniture(section)

    def _set_running_furniture(self, section):
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"dterrain  |  {self.short_title}")
        set_run_font(r, size=9, bold=True, color=MUTED)

        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"版本 {VERSION}  |  第 ")
        set_run_font(r, size=9, color=MUTED)
        add_page_field(p)
        r = p.add_run(" 页")
        set_run_font(r, size=9, color=MUTED)

    def cover(self, kicker):
        for _ in range(5):
            self.doc.add_paragraph().paragraph_format.space_after = Pt(10)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(18)
        r = p.add_run(kicker.upper())
        set_run_font(r, size=10, bold=True, color=BLUE)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(12)
        r = p.add_run(self.title)
        set_run_font(r, size=28, bold=True, color=NAVY)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(44)
        r = p.add_run(self.subtitle)
        set_run_font(r, size=14, color=DARK_BLUE)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(f"版本 {VERSION}  |  {TODAY}")
        set_run_font(r, size=11, bold=True, color=NAVY)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"适用对象：{self.audience}")
        set_run_font(r, size=10, italic=True, color=MUTED)
        self.doc.add_page_break()

    def h1(self, text):
        self._number_counter = 0
        return self.doc.add_paragraph(text, style="Heading 1")

    def h2(self, text):
        self._number_counter = 0
        return self.doc.add_paragraph(text, style="Heading 2")

    def h3(self, text):
        self._number_counter = 0
        return self.doc.add_paragraph(text, style="Heading 3")

    def para(self, text="", bold_prefix=None):
        p = self.doc.add_paragraph()
        if bold_prefix and text.startswith(bold_prefix):
            r = p.add_run(bold_prefix)
            set_run_font(r, bold=True)
            r = p.add_run(text[len(bold_prefix):])
            set_run_font(r)
        else:
            r = p.add_run(text)
            set_run_font(r)
        return p

    def bullet(self, text):
        p = self.doc.add_paragraph(style="List Bullet")
        r = p.add_run(text)
        set_run_font(r)
        return p

    def number(self, text):
        self._number_counter += 1
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.20)
        r = p.add_run(f"{self._number_counter}.  {text}")
        set_run_font(r)
        return p

    def callout(self, label, text, tone="blue"):
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.16)
        p.paragraph_format.right_indent = Inches(0.16)
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.2
        fill = {"blue": LIGHT_BLUE, "gray": CALLOUT,
                "gold": "FFF7E0", "red": "FDECEC"}[tone]
        set_paragraph_shading(p, fill)
        r = p.add_run(label + "  ")
        set_run_font(r, bold=True,
                     color={"blue": DARK_BLUE, "gray": NAVY,
                            "gold": GOLD, "red": RED}[tone])
        r = p.add_run(text)
        set_run_font(r)
        return p

    def code(self, text):
        p = self.doc.add_paragraph(style="Dterrain Code")
        set_paragraph_shading(p, LIGHT_GRAY)
        lines = text.strip("\n").split("\n")
        for index, line in enumerate(lines):
            if index:
                p.add_run().add_break()
            r = p.add_run(line)
            set_run_font(r, size=8.5, latin="Consolas",
                         east_asia="Microsoft YaHei")
        return p

    def table(self, headers, rows, widths_dxa, alignments=None):
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        keep_table_row_together(table.rows[0])
        for index, header in enumerate(headers):
            cell = table.rows[0].cells[index]
            shade_cell(cell, LIGHT_BLUE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.keep_with_next = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(header))
            set_run_font(r, size=9.5, bold=True, color=NAVY)
        set_repeat_table_header(table.rows[0])
        for row_values in rows:
            row = table.add_row()
            keep_table_row_together(row)
            cells = row.cells
            for index, value in enumerate(row_values):
                p = cells[index].paragraphs[0]
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                if alignments:
                    p.alignment = alignments[index]
                r = p.add_run(str(value))
                set_run_font(r, size=9.2)
        set_table_geometry(table, widths_dxa)
        after = self.doc.add_paragraph()
        after.paragraph_format.space_before = Pt(0)
        after.paragraph_format.space_after = Pt(2)
        return table

    def metadata(self, rows):
        self.table(["项目", "说明"], rows, [2700, 6660],
                   [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT])

    def new_page(self):
        self.doc.add_page_break()

    def save(self, path):
        path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.core_properties.title = self.title
        self.doc.core_properties.subject = self.subtitle
        self.doc.core_properties.author = "dterrain project"
        self.doc.core_properties.keywords = "Delaunay,TIN,CGAL,terrain,DLL"
        self.doc.core_properties.comments = "Generated from the dterrain source tree."
        self.doc.save(path)


def add_document_control(m: Manual, scope, navigation):
    m.h1("文档说明")
    m.metadata([
        ("文档版本", VERSION),
        ("发布日期", TODAY),
        ("适用平台", "Windows x64；C++17；MinGW 或 Visual Studio"),
        ("适用工程", "dterrain 动态 2.5D Delaunay TIN 动态库及演示程序"),
        ("内容范围", scope),
    ])
    m.h2("阅读导航")
    for item in navigation:
        m.bullet(item)
    m.callout("版本边界", "本手册对应 dterrain 0.27.0：新增世界 XY 视口到 GRID 最小源节点窗口的仿射裁剪接口，并由 GUI 在缩放、平移、拉框放大和全图适屏后，仅重读当前视口的最多 512×512 局部 LOD。接口支持旋转、剪切和负像素高仿射；完整 GRID 仍参与统计、等高线、分析与导出。窗口概览、多边形裁剪、显式重采样、解析挖填方、全幅地形专题、剖面、GDAL 与约束事务保持兼容。坐标重投影、磁盘流式瓦片、持久化多级金字塔、网格单元解析边界裁剪、局部 CDT 和 GPU LOD 属于后续阶段。", "gold")


def build_developer_manual():
    m = Manual(
        "dterrain 动态地形三角网 DLL\n开发与使用手册",
        "DLL 开发与使用手册",
        "面向测绘地形建模的 TIN、GRID、等高线转换、视口自适应 GRID LOD、动态编辑与空间查询",
        "C/C++ 集成人员、算法研究人员、测试与运维人员",
    )
    m.cover("DEVELOPER REFERENCE")
    add_document_control(
        m,
        "架构、构建部署、稳定 C ABI、世界视口 GRID 裁剪与自适应 LOD、GRID 窗口概览、GRID 多边形裁剪/掩膜、GRID 显式重采样、双 GRID 精确挖填方、局部与全幅坡度/坡向/阴影分析、约束 Delaunay、GRID/等高线转换、GDAL 格式交换、异步任务、12 个兼容接口、文件格式、性能、线程安全、故障排查。",
        [
            "首次集成：重点阅读第 3、4、7 章。",
            "地形转换：重点阅读第 4.6 章和第 6 章。",
            "兼容既有系统：重点阅读第 5 章的 12 个 Legacy 接口。",
            "大数据量使用：重点阅读第 8 章的内存与查询策略。",
            "文件交换：重点阅读第 4.7 章和第 6 章。",
            "现状/设计面土方：重点阅读第 4.11 节。",
            "错位 GRID 对齐：重点阅读第 4.12 节。",
            "任意区域 GRID 裁剪：重点阅读第 4.13 节。",
            "大 GRID 概览与 LOD：重点阅读第 4.14 节。",
            "视口自适应 GRID LOD：重点阅读第 4.15 节。",
        ],
    )

    m.h1("1 产品概述")
    m.para("dterrain 是一个 C++17 动态库，在 XY 平面构建 Delaunay 三角网，并把 Z 作为顶点高程属性保存。它适用于局部坐标或投影坐标下的测绘地形建模、算法研究、动态编辑演示和既有系统集成。")
    m.h2("1.1 已实现能力")
    for text in (
        "从内存点数组或 XYZ/CSV 文本批量构建 Delaunay TIN。",
        "动态插入、按 XY 最近点删除、按稳定 ID 删除以及高程更新。",
        "返回编辑前后受影响三角形、边界边、删除边与新增边。",
        "最近顶点、点定位、矩形范围三角形查询、统计与完整校验。",
        "GUI 两点定义任意地形剖面，固定 CDT/TIN/GRID 数据源采样、统计并导出 CSV。",
        "GUI 逐点圈定简单多边形，计算平面量并数值估算地表面积和水平基准面挖填方。",
        "普通 TIN、约束 CDT 和仿射 GRID 统一返回局部高程梯度、坡度、下坡坡向、单位法向与支撑单元。",
        "从任意仿射 GRID 派生全幅坡度角、下坡坡向角或 0～255 分析阴影，并继承 CRS/NoData。",
        "匹配现状/设计 GRID 逐单元解析积分挖填方，支持差值 GRID、并行、进度和取消。",
        "同 CRS GRID 按完整六参数仿射显式对齐，支持最近邻、双线性、NoData 策略、并行与取消。",
        "GRID 按任意简单多边形作保持范围掩膜、紧凑裁剪或反向掩膜，支持旋转/剪切仿射、并行与取消。",
        "GRID 指定窗口直接读取 caller-owned 概览，支持平均、最近邻、最小值、最大值、NoData、精确统计与确定性并行。",
        "世界 XY 视口经完整六参数仿射反算与多边形裁剪得到最小 GRID 节点窗口，可直接组合窗口概览生成局部 LOD。",
        "DTIN 二进制保存加载，以及 DTMESH 可读文本三角网交换。",
        "双精度 GRID、仿射节点坐标、NoData、窗口读写和 DGRID 文本往返。",
        "TIN→GRID、GRID→TIN、TIN/GRID→等高线，以及等高线→TIN/GRID 近似重建。",
        "耗时转换的异步任务、进度、等待、协作取消和结果提取。",
        "TIN/GRID/等高线 CRS WKT 元数据，以及可选 GeoTIFF/COG/GeoPackage 交换。",
        "独立 CDT 句柄、断裂线、外边界、孔洞、域内查询、约束增删与 DCDT 文本往返。",
        "保持约束 ID 和类型的原子几何更新；可选返回更新前后的完整域差异。",
        "约束顶点引用查询、共享顶点默认保护和显式单约束脱离删除。",
        "原子批量约束事务，一次提交多条新增、更新和删除并统一回滚。",
        "普通 TIN→CDT、CDT 域内高程采样，以及保留外边界/孔洞的 GRID 和等高线派生。",
        "推荐的稳定 C ABI 与原需求 12 个 C++ 接口兼容层。",
    ):
        m.bullet(text)
    m.h2("1.2 核心语义")
    m.table(
        ["主题", "规则"],
        [
            ("数据模型", "2.5D：Delaunay 判定只使用 XY；Z 是顶点高程属性。"),
            ("重复点", "相同 XY 被视为重复，返回 DT_E_DUPLICATE_XY。"),
            ("最近距离", "最近顶点和删除最近点都使用 XY 平面欧氏距离。"),
            ("高程更新", "dt_update_vertex_z() 只改变 Z，不改变网格拓扑。"),
            ("事务性", "批量建网和文件加载失败时保留原三角网。"),
        ("转换边界", "等高线反推只能由有限折点近似恢复表面；普通 TIN 不强制保留原折线边。"),
        ],
        [2200, 7160],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
    )
    m.callout("适用边界", "GRID→TIN 和等高线→TIN 都重建普通 Delaunay；严格折线、边界和孔洞应使用独立 dt_cdt_handle。v0.24 单项约束编辑仍采用候选网全量重建；批量事务把 N 次重建降为 1 次，但尚未达到百万点实时局部 CDT 编辑。GUI 多边形土方针对水平基准面并使用固定预算数值积分；双 GRID 接口要求节点对齐，错位表面可先显式重采样。所有土方结果均需结合项目规范复核后用于工程结算。", "gold")

    m.h1("2 架构与数据模型")
    m.h2("2.1 分层结构")
    m.code("""调用系统 / GUI
        |
稳定 C ABI + CDT API + Terrain API + Task API + GDAL API + Legacy 兼容层
        |
上下文、错误处理、结果句柄
        |
CGAL Delaunay hierarchy ---- Boost R-tree ---- 独立 Constrained Delaunay
        |                         |
顶点稳定 ID / Z              范围相交查询
        |
DTIN / DTMESH / XYZ / DGRID / DCONTOUR / DCDT 持久化
        |
GRID <---- 转换引擎 ----> TIN <----> 等高线
  \\------------- 可选 GDAL Adapter -------------/""")
    m.para("CGAL 顶点几何只保存 XY，自定义顶点信息保存 uint64_t 稳定 ID 和 Z。Boost R-tree 保存有限三角形的 XY 包围盒与面句柄。局部编辑时先移除旧面索引项，再加入新面，因此不需要每次重建全局范围索引。")
    m.h2("2.2 关键技术选择")
    m.table(
        ["组件", "用途", "工程影响"],
        [
            ("CGAL EPICK", "精确谓词、非精确构造", "提高方向和空圆判定鲁棒性。"),
            ("Delaunay_triangulation_2", "二维 Delaunay 拓扑", "Z 不进入三角剖分判定。"),
            ("Triangulation_hierarchy_2", "分层点定位", "有利于动态插入和空间定位。"),
            ("Boost R-tree", "面包围盒索引", "加速当前视口和矩形范围查询。"),
            ("Terrain Core", "GRID 与等高线转换", "不向 DLL 外暴露 C++ 容器和 CGAL 类型。"),
            ("Task Runtime", "后台转换任务", "保持源对象生命周期并提供协作取消。"),
            ("Constrained Delaunay", "断裂线与域裁剪", "独立句柄保持普通 TIN ABI 和性能语义。"),
        ],
        [2400, 2500, 4460],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )

    m.h1("3 获取、构建与部署")
    m.h2("3.1 依赖")
    m.table(
        ["依赖", "最低/建议版本", "说明"],
        [
            ("CMake", "3.24+", "生成构建系统与安装目录。"),
            ("C++ 编译器", "C++17", "MinGW-w64 GCC 或 Visual Studio 2022 x64。"),
            ("CGAL", "6.x", "核心 Delaunay 后端。"),
            ("Boost", "随 CGAL", "R-tree 与 CGAL 依赖。"),
            ("GMP/MPFR", "随 vcpkg", "CGAL 数值依赖。"),
        ],
        [1800, 2200, 5360],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
    )
    m.h2("3.2 MinGW + vcpkg 构建")
    m.code("""cmake -S . -B build -G "MinGW Makefiles" ^
  -DCMAKE_TOOLCHAIN_FILE=%VCPKG_ROOT%/scripts/buildsystems/vcpkg.cmake ^
  -DVCPKG_TARGET_TRIPLET=x64-mingw-dynamic ^
  -DCMAKE_BUILD_TYPE=Release ^
  -DDT_BUILD_TESTS=ON -DDT_BUILD_EXAMPLES=ON -DDT_BUILD_GUI=ON

cmake --build build --parallel 4
ctest --test-dir build --output-on-failure
cmake --install build --prefix dist""")
    m.h2("3.3 Visual Studio 构建")
    m.para("将生成器改为 Visual Studio 2022，并使用 x64-windows triplet。调用程序、DLL 和第三方运行库必须使用兼容的架构与运行时设置，避免混用 x86/x64 或不同 CRT。")
    m.h2("3.4 分发清单")
    m.table(
        ["文件", "是否必须", "用途"],
        [
            ("dterrain.dll", "是", "动态库实现。"),
            ("dt_api.h", "开发期", "推荐的稳定 C ABI。"),
            ("dt_cdt_api.h", "按需", "约束 Delaunay、边界与孔洞接口。"),
            ("dt_terrain_api.h", "开发期", "GRID、等高线和同步转换接口。"),
            ("dt_task_api.h", "按需", "异步转换任务接口。"),
            ("dt_legacy.hpp", "按需", "12 个原始接口兼容声明。"),
            ("libgmp-10.dll", "MinGW 包", "CGAL/GMP 运行库。"),
            ("libwinpthread-1.dll", "MinGW 包", "线程运行库。"),
            ("libdterrain.dll.a / .lib", "链接期", "导入库。"),
        ],
        [3300, 1700, 4360],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
    )
    m.callout("部署建议", "优先直接复制 dist/bin、dist/include 和 dist/lib 的对应内容。不要只复制 dterrain.dll 而遗漏 MinGW 运行库。", "blue")

    m.h1("4 稳定 C ABI 集成")
    m.h2("4.1 生命周期")
    m.code("""#include "dt_api.h"

dt_handle mesh = nullptr;
if (dt_create(nullptr, &mesh) != DT_OK) {
    // 调用 dt_get_last_error() 获取原因
    return;
}

// 使用 mesh
dt_destroy(mesh);""")
    m.para("每个 dt_handle 表示一个独立网格上下文。接口不会把 C++ 异常传播到 DLL 外部。dt_destroy() 接受有效句柄并释放其全部网格和索引资源。")
    m.h2("4.2 批量建网")
    m.code("""dt_point3 points[] = {
    {0.0, 0.0, 10.0},
    {100.0, 0.0, 12.5},
    {0.0, 100.0, 20.0},
    {100.0, 100.0, 28.0}
};
dt_vertex_id ids[4]{};
dt_status s = dt_build(mesh, points, 4, ids);""")
    m.para("output_ids 可以为 NULL。成功后整体替换当前网格；输入非法、重复 XY 或内存不足时返回错误并保留旧网格。百万级以上点集应优先使用 dt_build() 或 dt_import_points_text()，不要逐点调用插入接口。")
    m.h2("4.3 动态编辑和影响结果")
    m.code("""dt_point3 p{50.0, 40.0, 18.0};
dt_vertex_id id = 0;
dt_edit_result result = nullptr;

if (dt_insert_point(mesh, &p, &id, &result) == DT_OK) {
    dt_edit_result_view view{};
    dt_edit_result_get_view(result, &view);
    // 消费 removed_triangles / boundary_edges / added_edges 等
}
dt_release_edit_result(result);""")
    m.table(
        ["结果字段", "含义", "有效期"],
        [
            ("removed_triangles", "编辑前被移除的有限三角形", "结果句柄释放前"),
            ("added_triangles", "编辑后新增的有限三角形", "结果句柄释放前"),
            ("boundary_edges", "局部影响区边界", "结果句柄释放前"),
            ("removed_edges", "旧线框中被删除的边", "结果句柄释放前"),
            ("added_edges", "新线框中新增的边", "结果句柄释放前"),
            ("affected_vertex_id", "插入或删除所影响的稳定顶点 ID", "值复制后长期有效"),
        ],
        [2700, 4300, 2360],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER],
    )

    m.h2("4.4 查询")
    m.table(
        ["接口", "输入", "返回"],
        [
            ("dt_find_nearest_vertex_xy", "查询 XYZ；只使用 XY", "最近顶点坐标和 ID"),
            ("dt_locate_point_xy", "查询 XYZ；只使用 XY", "面、边、顶点或凸包外分类"),
            ("dt_query_triangles", "闭合 XY 矩形", "与矩形相交的有限三角形结果句柄"),
            ("dt_get_statistics", "网格句柄", "顶点数、面数、维度、范围、generation"),
            ("dt_validate", "网格句柄与详细级别", "CGAL 拓扑与 Delaunay 完整校验"),
        ],
        [3300, 2700, 3360],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    m.code("""dt_bounds2 box{1000.0, 2000.0, 1200.0, 2200.0};
dt_query_result result = nullptr;
if (dt_query_triangles(mesh, &box, &result) == DT_OK) {
    dt_query_result_view view{};
    dt_query_result_get_view(result, &view);
    // view.triangles[0 .. triangle_count-1]
}
dt_release_query_result(result);""")
    m.callout("大范围查询", "范围结果一次性保存在结果对象中。千万级全图可能产生非常大的三角形数组；显示程序应按当前视口查询，必要时再在渲染层抽样。", "gold")

    m.h2("4.5 文件接口")
    m.table(
        ["接口", "格式", "用途"],
        [
            ("dt_import_points_text", "XYZ/TXT/CSV", "读取散点并立即自动构网。"),
            ("dt_save_mesh_text", "DTMESH 1", "保存顶点、稳定 ID 和显式三角形。"),
            ("dt_load_mesh_text", "DTMESH 1", "重建并逐面验证 Delaunay 拓扑。"),
            ("dt_save / dt_load", "DTIN v1", "紧凑二进制点集保存加载。"),
            ("dt_grid_save/load_text", "DGRID 1", "规则高程节点文本往返。"),
            ("dt_contours_save/load_text", "DCONTOUR 1", "等高折线文本往返。"),
            ("dt_cdt_save/load_text", "DCDT 1", "散点、约束和 CRS 文本往返。"),
        ],
        [3200, 1800, 4360],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
    )
    m.code("""dt_bounds2 loaded{};
dt_status s = dt_import_points_text(
    mesh, "sample_data/sample_points.xyz", &loaded);
if (s != DT_OK) {
    char error[512]{};
    dt_get_last_error(error, sizeof(error), nullptr);
}""")

    m.h2("4.6 GRID、等高线与异步转换")
    m.para("新增接口分别位于 dt_terrain_api.h 和 dt_task_api.h。GRID 节点采用六参数仿射变换定位；TIN→GRID 在覆盖三角面上进行分片线性插值，凸包外写入 NoData。")
    m.code("""dt_tin_to_grid_options options{};
options.struct_size = sizeof(options);
options.width = 1001;
options.height = 1001;
options.geo_transform[0] = xmin;
options.geo_transform[1] = (xmax - xmin) / 1000.0;
options.geo_transform[3] = ymin;
options.geo_transform[5] = (ymax - ymin) / 1000.0;
options.nodata_value = -9999.0;

dt_grid_handle grid = nullptr;
dt_grid_from_tin(mesh, &options, &grid);
dt_grid_destroy(grid);""")
    m.para("v0.12 可把等高线的 LINE elevation 作为权威高程，以折点或按最大段长加密后的样本重建普通 TIN，并可直接插值 GRID。重复 XY 高程一致时合并，冲突时整次转换失败且原 TIN 不变。")
    m.code("""dt_contours_to_tin_options sampling{};
sampling.struct_size = sizeof(sampling);
sampling.maximum_segment_length = 5.0; // 0: 只用原折点
sampling.merge_tolerance = 1.0e-8;      // 0: 只合并完全相同 XY

dt_tin_from_contours(contours, &sampling, mesh);
dt_grid_from_contours(contours, &sampling, &options, &grid);""")
    m.callout("反向转换边界", "等高线之间的极值和曲面细节已经丢失，不能唯一恢复。输出是普通 Delaunay，原等高线段不保证成为网边；硬折线应使用 CDT。", "gold")
    m.table(
        ["接口组", "主要功能", "释放方式"],
        [
            ("dt_grid_*", "GRID 创建、信息、窗口读写、文本往返", "dt_grid_destroy"),
            ("dt_*_from_*", "TIN/GRID/等高线同步转换与近似重建", "按输出句柄释放"),
            ("dt_task_*", "异步启动、进度、等待、取消、结果提取", "dt_task_destroy"),
            ("dt_contours_*", "等高线信息、逐线视图和文本往返", "dt_contours_destroy"),
        ],
        [2300, 4700, 2360],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER],
    )
    m.callout("NoData", "GRID→TIN 遇到 NoData 默认返回 DT_E_UNSUPPORTED。DT_GRID_TO_TIN_ALLOW_NODATA_BRIDGING 只适合调用方明确接受跨空洞三角化的场景。", "gold")
    m.code("""dt_task_handle task = nullptr;
dt_grid_from_tin_async(mesh, &options, &task);
int32_t completed = 0;
dt_task_wait(task, UINT32_MAX, &completed);

dt_task_info info{};
dt_task_get_info(task, &info);
if (info.state == DT_TASK_SUCCEEDED) {
    dt_grid_handle result = nullptr;
    dt_task_get_grid_result(task, &result);
    dt_grid_destroy(result);
}
dt_task_destroy(task);""")

    m.h2("4.7 CRS 与 GDAL 格式交换")
    m.para("dt_gdal_api.h 是可选适配层。默认 DT_WITH_GDAL=OFF 时函数仍然导出，但返回 DT_E_UNSUPPORTED；启用后可导入导出 GeoTIFF/COG GRID 和 GeoPackage 等高线。TIN、GRID 与等高线的 CRS 以 UTF-8 WKT 保存并在转换时传播，当前不执行坐标重投影。")
    m.code("""#include "dt_gdal_api.h"

dt_gdal_raster_save_options save{};
save.struct_size = sizeof(save);
save.driver_name = "COG";
const char* co[] = {"COMPRESS=DEFLATE", nullptr};
save.creation_options = co;
dt_grid_save_gdal_raster(grid, "terrain.tif", &save);

dt_grid_handle loaded = nullptr;
dt_grid_load_gdal_raster("terrain.tif", nullptr, &loaded);
dt_grid_destroy(loaded);""")
    m.table(
        ["对象", "格式/驱动", "保真内容"],
        [
            ("GRID", "GTiff / COG", "double 高程、NoData、仿射变换、CRS WKT"),
            ("等高线", "GPKG", "LineStringZ、elevation、closed、CRS WKT"),
            ("坐标换算", "GDAL ↔ GRID", "像元角点与节点中心自动进行半像元偏移"),
        ],
        [2000, 2500, 4860],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    m.callout("部署", "启用 GDAL 的便携目录必须携带构建目录复制出的 GDAL、PROJ、SQLite、GeoTIFF 等运行时 DLL，以及 DLL 同级 share/proj/proj.db。COG 使用 CreateCopy，会比普通 GTiff 需要更多临时内存/存储。", "gold")
    m.callout("GUI 集成", "v0.13 的演示程序启动时探测 GTiff、COG、GPKG 驱动，按能力启用五个数据交换菜单。栅格导入只替换 GRID，GeoPackage 导入只替换等高线，普通 TIN、CDT 和无关图层保持不变。", "blue")

    m.h2("4.8 约束 Delaunay")
    m.para("dt_cdt_api.h 提供独立 dt_cdt_handle。基础点与普通 TIN 相互独立；断裂线只强制成边，外边界和孔洞通过奇偶嵌套层级决定有效域。没有外边界时全部有限面均为有效域。")
    m.code("""#include "dt_cdt_api.h"

dt_cdt_handle cdt = nullptr;
dt_cdt_create(nullptr, &cdt);
dt_cdt_build_from_tin(cdt, tin);
dt_constraint_id boundary_id = 0;
dt_cdt_add_constraint(cdt, DT_CONSTRAINT_OUTER_BOUNDARY, 0,
                      boundary, boundary_count, &boundary_id);
dt_cdt_add_constraint(cdt, DT_CONSTRAINT_HOLE_BOUNDARY, 0,
                      hole, hole_count, nullptr);

// 保持 boundary_id 和 OUTER_BOUNDARY 类型，原子替换几何。
// 不需要影响区时将最后一个参数设为 nullptr，可避免完整差异计算。
dt_edit_result effect = nullptr;
if (dt_cdt_update_constraint(cdt, boundary_id, 0,
                             moved_boundary, moved_boundary_count,
                             &effect) == DT_OK) {
    dt_edit_result_view view{sizeof(dt_edit_result_view)};
    dt_edit_result_get_view(effect, &view);
    // view 中包含删除面、新增面、边界边、删除边和新增边。
    dt_release_edit_result(effect);
}""")
    m.code("""dt_cdt_vertex_usage usage{};
dt_cdt_get_constraint_vertex_usage(cdt, boundary_id, 1, &usage);
uint32_t remove_flags = usage.constraint_count > 1
    ? DT_CDT_REMOVE_VERTEX_ALLOW_SHARED_DETACH : 0;
dt_cdt_remove_constraint_vertex(cdt, boundary_id, 1,
                                remove_flags, nullptr);

dt_cdt_constraint_edit edits[2]{};
edits[0].struct_size = sizeof(edits[0]);
edits[0].operation = DT_CDT_EDIT_UPDATE;
edits[0].constraint_id = boundary_id;
edits[0].points = moved_boundary;
edits[0].point_count = moved_boundary_count;
edits[1].struct_size = sizeof(edits[1]);
edits[1].operation = DT_CDT_EDIT_ADD;
edits[1].kind = DT_CONSTRAINT_BREAKLINE;
edits[1].points = ridge;
edits[1].point_count = ridge_count;
dt_constraint_id result_ids[2]{};
dt_cdt_apply_constraint_edits(cdt, edits, 2, result_ids, nullptr);

dt_grid_from_cdt(cdt, &grid_options, &grid);
dt_contours_from_cdt(cdt, &contour_options, &contours);
dt_cdt_save_text(cdt, "terrain.dcdt");
dt_cdt_destroy(cdt);""")
    m.table(
        ["类型", "作用", "闭合规则"],
        [
            ("BREAKLINE", "强制三角网沿折线分边，不改变域内外", "可开可闭"),
            ("OUTER_BOUNDARY", "地形有效域外边界", "自动闭合"),
            ("HOLE_BOUNDARY", "从外边界域中排除孔洞", "自动闭合；必须已有外边界"),
        ],
        [2600, 4300, 2460],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    m.para("dt_cdt_get_constraint_vertex_usage() 报告指定点被多少条约束、多少个点序列位置引用，以及是否也是基础地形点。dt_cdt_remove_constraint_vertex() 默认拒绝共享顶点；显式传 DT_CDT_REMOVE_VERTEX_ALLOW_SHARED_DETACH 时只从目标约束脱离，其他约束和基础点保持不变。删除后仍统一校验最小点数和拓扑，失败时 generation 不变。")
    m.para("dt_cdt_apply_constraint_edits() 按数组顺序组合 ADD、UPDATE 和 REMOVE。ADD 分配新稳定 ID，UPDATE 保留原类型，REMOVE 不携带几何。任一参数、ID 或最终拓扑失败时整批回滚且不消耗 ID；成功时 generation 只增加一次。output_constraint_ids 和整批 output_effect 均可为空。")
    m.para("dt_cdt_sample_height_xy() 只在有效域内插值；外边界以外和孔洞内部返回 DT_E_NOT_FOUND。dt_grid_from_cdt() 将这些位置写为 NoData，dt_contours_from_cdt() 只遍历域内三角形并继承 CRS。更新和顶点删除的 output_effect 均可为 nullptr。")
    m.callout("性能与交叉", "v0.27 单项约束编辑仍完整重建候选 CDT；批量事务只重建一次。100,000 基础点添加 12 条断裂线的本机对比为逐条 3.259 s、批量 0.528 s，约 6.17 倍。请求 output_effect 还会计算完整域差异。未分段交叉仍须先加入共享顶点。", "gold")

    m.h2("4.9 坡度、坡向与地形法向")
    m.para("v0.16 由普通 TIN、约束 CDT 和仿射 GRID 共用 dt_surface_analysis 结果结构。三个入口分别是 dt_analyze_tin_surface_xy()、dt_cdt_analyze_surface_xy() 和 dt_grid_analyze_surface_xy()；查询 Z 被忽略，成功结果给出插值 Z、世界 XY 梯度、坡度角、下坡坡向、向上单位法向和实际支撑点。接口只写调用方结构，不分配跨 DLL 内存。")
    m.code("""dt_surface_analysis a{};
dt_point3 q{500000.0, 3200000.0, 0.0};
dt_status s = dt_analyze_tin_surface_xy(tin, &q, &a);
if (s == DT_OK) {
    // a.slope_degrees: 相对水平面的坡度角
    // a.aspect_degrees: 从 +Y 北向顺时针的下坡坡向
    // a.normal_*: 指向上方的单位法向
}""")
    m.table(
        ["字段/标志", "约定"],
        [
            ("dz_dx / dz_dy", "世界 XY 坐标中的局部高程梯度。"),
            ("slope_degrees", "atan(hypot(dz_dx,dz_dy))，水平面为 0°。"),
            ("aspect_degrees", "最大下降方向；+Y 为北，顺时针 0～360°。"),
            ("DT_SURFACE_ASPECT_UNDEFINED", "水平面没有唯一坡向，此时角度字段为 0。"),
            ("QUERY_ON_EDGE / QUERY_ON_VERTEX", "TIN/CDT 在多个邻面处选择一个有效支撑面。"),
            ("DT_SURFACE_BILINEAR", "GRID 使用局部 2×2 节点双线性导数。"),
        ],
        [3100, 6260],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    m.para("TIN 与 CDT 在支撑三角面内使用解析平面梯度；CDT 外域和孔洞返回 DT_E_NOT_FOUND。GRID 先反算完整六参数仿射坐标，在 2×2 单元求双线性导数，再转换到世界 XY；四个支撑节点任一为 NoData 时拒绝分析。成功时 DLL 在 struct_size 中写入本版本结构大小，接口不分配需要调用方释放的内存。")
    m.callout("边和顶点", "分片线性 TIN/CDT 在断裂边两侧可有不同坡度。查询恰落在边或顶点时，本版本确定性选择一个有限有效邻面并返回位置标志，不做跨断裂线平滑平均。", "gold")

    m.h2("4.10 全幅 GRID 地形专题分析")
    m.para("dt_grid_derive_terrain() 从源高程 GRID 创建同尺寸、同六参数仿射变换和同 CRS 的标准 dt_grid_handle；源句柄保持不变，输出可继续使用窗口读取、DGRID 保存或 GDAL GeoTIFF/COG 导出。v0.21 接入统一异步任务，v0.22 在相同同步/异步入口下增加分块多线程。输出始终具有 NoData。")
    m.code("""dt_grid_terrain_options o{};
o.struct_size = sizeof(o);
o.kind = DT_GRID_TERRAIN_SLOPE_DEGREES;
o.z_factor = 1.0;
o.output_nodata_value = -9999.0;
o.worker_count = 0;    // 自动线程，内部最多 32
o.tile_row_count = 0;  // 默认 64 行一块

dt_grid_handle slope = nullptr;
dt_status s = dt_grid_derive_terrain(elevation, &o, &slope);
if (s == DT_OK) {
    dt_grid_save_text(slope, "slope.dgrid");
    dt_grid_destroy(slope);
}""")
    m.code("""dt_task_handle task = nullptr;
dt_grid_derive_terrain_async(elevation, &o, &task);
for (;;) {
    int32_t done = 0;
    dt_task_wait(task, 50, &done);
    dt_task_info info{};
    dt_task_get_info(task, &info);  // progress: 0..1
    if (done) break;
    // 用户取消：dt_task_request_cancel(task);
}
dt_grid_handle result = nullptr;
if (dt_task_get_grid_result(task, &result) == DT_OK) {
    // 使用 result，最终 dt_grid_destroy(result)
}
dt_task_destroy(task);""")
    m.table(
        ["kind", "输出", "特殊语义"],
        [
            ("SLOPE_DEGREES", "坡度角 0～90°", "水平面为 0°"),
            ("ASPECT_DEGREES", "下坡方位角 0～360°", "水平面无唯一坡向，写 NoData"),
            ("HILLSHADE", "分析阴影 0～255", "默认光源方位/高度 315°/45°"),
        ],
        [2600, 3000, 3760],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    m.para("内部节点采用中心差分，边界采用单边差分；列/行导数通过完整仿射雅可比转换为世界 XY 梯度。中心或所需四邻域含 NoData 时结果节点写 NoData，在空洞边缘形成一节点安全带。z_factor 为零时选择 1.0，否则必须有限且大于零。output_nodata_value 为零时自动选择 NaN，避免与有效 0°/0 灰度冲突。输出句柄必须由 dt_grid_destroy() 释放。")
    m.table(
        ["字段", "零值", "有效非零值"],
        [
            ("worker_count", "自动，最多 32", "1 单线程；2～64 请求线程数"),
            ("tile_row_count", "64 行", "1～1048576 行"),
        ],
        [2500, 2800, 4060],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    m.callout("并行、异步与内存", "选项仍为 80 字节，字段复用原 ABI 预留区。工作线程只写互不重叠的输出行，调用线程串行执行进度/取消回调；取消状态为 DT_TASK_CANCELLED/DT_E_CANCELLED，不发布半成品。无需第二份全幅临时数组，除源 GRID 外主要增加一个同尺寸 double 输出 GRID。z_factor 用于高程和平面单位换算；XY 为米、Z 为毫米时使用 0.001。", "blue")
    m.code("""# DT_BUILD_BENCHMARKS=ON
dterrain_terrain_benchmark.exe 4096 4096 0 64
# 输出串行/并行秒数、加速比和抽样校验和是否一致""")
    m.callout("本机基准", "4096×4096 坡度 GRID：单线程 0.233 s，自动并行 0.0678 s，约 3.44×，64 点抽样校验和一致。收益受 CPU、内存带宽、尺寸、块高和专题类型影响，不是固定承诺。", "gold")

    m.h2("4.11 现状/设计双 GRID 精确挖填方")
    m.para("dt_grid_compare_earthwork() 比较节点数、完整六参数仿射变换和 CRS 一致的现状 GRID 与设计 GRID。高差定义为 existing×existing_z_factor − design×design_z_factor：正值计挖方，负值的绝对体积计填方，净方为挖方减填方。输出统计按世界 XY 面积计算，因此旋转、剪切或负像元高的仿射 GRID 同样适用。")
    m.code("""dt_grid_earthwork_options o{};
o.struct_size = sizeof(o);
o.flags = DT_GRID_EARTHWORK_OUTPUT_DIFFERENCE_GRID;
o.worker_count = 0;   // 自动最多 32；显式最多 64
o.tile_row_count = 0; // 默认 64 个单元行
o.existing_z_factor = 1.0;
o.design_z_factor = 1.0;
o.output_nodata_value = -9999.0;

dt_grid_earthwork_result r{};
r.struct_size = sizeof(r);
dt_grid_handle difference = nullptr;
dt_status s = dt_grid_compare_earthwork(existing, design, &o,
                                        &r, &difference);
if (s == DT_OK) {
    // 使用 r.cut_volume / r.fill_volume / r.net_volume
    dt_grid_save_text(difference, "difference.dgrid");
    dt_grid_destroy(difference);
}""")
    m.table(
        ["结果", "定义", "说明"],
        [
            ("cut/fill/net_volume", "挖方、填方、挖方−填方", "按每个网格单元的两个线性三角面解析积分"),
            ("total/valid_area", "总平面面积、有效平面面积", "coverage_ratio = valid_area / total_area"),
            ("min/max/mean_difference", "节点高差范围与面积加权均值", "符号遵循现状减设计"),
            ("rms_difference", "高差均方根", "使用线性三角面上 d² 的解析积分"),
            ("difference GRID", "逐节点现状减设计", "可选；保持现状 GRID 的仿射和 CRS"),
        ],
        [2300, 3300, 3760],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    m.para("每个四节点单元固定按左上—右下对角线剖分。高差全同号时直接积分；跨越零高差线时，在三角形边上解析求交并分别积分挖、填子多边形，不依赖采样密度。默认任一角点为 NoData 就跳过整个单元；设置 DT_GRID_EARTHWORK_ALLOW_PARTIAL_CELLS 后，可独立保留四角均有效的半单元三角形。")
    m.code("""dt_task_handle task = nullptr;
dt_grid_compare_earthwork_async(existing, design, &o, &task);
// 用 dt_task_get_info() 读取 0..1 进度；需要时请求取消
dt_grid_earthwork_result r{};
r.struct_size = sizeof(r);
dt_grid_handle difference = nullptr;
if (dt_task_get_earthwork_result(task, &r, &difference) == DT_OK) {
    // difference 为调用方拥有的新句柄
    dt_grid_destroy(difference);
}
dt_task_destroy(task);""")
    m.callout("校验与所有权", "两幅 GRID 的行列数、六参数仿射和 CRS WKT 必须匹配；源句柄在计算期间不得由其他线程修改。同步/异步取消或失败均不发布半成品。options/result 的 ABI 大小分别为 64/112 字节；返回的差值 GRID 由调用方 dt_grid_destroy()。", "blue")
    m.callout("本机基准", "4096×4096 双 GRID：单线程 0.439 s，自动并行 0.0381 s，约 11.50×；串并行体积差为 0。结果受 CPU、内存带宽、数据分布和块高影响，不是固定承诺。", "gold")

    m.h2("4.12 GRID 显式重采样与设计面对齐")
    m.para("dt_grid_resample_like() 把 source_grid 显式采样到 reference_grid 的行列数和完整六参数仿射几何上，输出继承参考 GRID 的 CRS。两者 CRS WKT 必须完全一致；接口绝不隐式重投影。该设计把坐标变换、插值和土方积分拆成可检查的步骤：先对齐设计面，再调用 dt_grid_compare_earthwork()。")
    m.code("""dt_grid_resample_options o{};
o.struct_size = sizeof(o);
o.method = DT_GRID_RESAMPLE_BILINEAR; // 0 也表示双线性
o.flags = 0;                          // 严格 NoData
o.worker_count = 0;                   // 自动最多 32；显式最多 64
o.tile_row_count = 64;
o.output_nodata_value = -9999.0;

dt_grid_handle aligned = nullptr;
dt_status s = dt_grid_resample_like(design, existing, &o, &aligned);
if (s == DT_OK) {
    // aligned 与 existing 节点严格对齐，可继续作双 GRID 土方
    dt_grid_destroy(aligned);
}""")
    m.table(
        ["方法/标志", "行为", "适用场景"],
        [
            ("NEAREST", "按源 GRID 最近节点取值", "分类值、快速预览或不希望平滑"),
            ("BILINEAR", "在源 GRID 的 2×2 节点上双线性插值", "连续高程设计面，GUI 默认推荐"),
            ("严格 NoData", "四个支撑点任一无效即输出 NoData", "避免跨越空洞插值；默认"),
            ("RENORMALIZE_NODATA", "忽略无效支撑点并归一化剩余权重", "明确接受空洞边缘外推时使用"),
        ],
        [2600, 3740, 3020],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    m.para("实现先把参考节点的世界 XY 仿射映射与源 GRID 逆仿射组合成列、行增量，因此旋转、剪切和负像元高均可正确处理。落在源节点包络以外的位置写输出 NoData。双线性模式要求源 GRID 至少 2×2；最近邻可处理单行或单列。output_nodata_value 为零时选择 NaN。")
    m.code("""dt_task_handle task = nullptr;
dt_grid_resample_like_async(design, existing, &o, &task);
// dt_task_get_info() 读取进度；Esc 等 UI 动作可请求协作取消
dt_grid_handle aligned = nullptr;
if (dt_task_get_grid_result(task, &aligned) == DT_OK) {
    // aligned 是调用方拥有的新句柄
    dt_grid_destroy(aligned);
}
dt_task_destroy(task);""")
    m.callout("并行、取消与所有权", "options 固定为 64 字节。工作线程按输出行块写互不重叠的区域，只有协调线程调用进度/取消回调；同步或异步失败、取消均不发布半成品。计算期间不得修改源和参考句柄，成功结果由调用方 dt_grid_destroy()。", "blue")
    m.callout("本机基准", "4096×4096 双线性重采样：单线程 0.243 s，自动并行 0.0538 s，约 4.53×，串并行抽样校验和误差为 0。收益受 CPU、内存带宽、仿射关系和块高影响，不是固定承诺。", "gold")

    m.h2("4.13 GRID 任意多边形裁剪与掩膜")
    m.para("dt_grid_clip_polygon() 用多边形的 XY 投影选择 GRID 节点。输入点的 Z 被忽略，多边形可省略重复闭合点，接口会隐式闭合。边界节点按内部处理；采用偶奇规则，因此结果对顺、逆时针一致。输出继承源 GRID 的 CRS 和 NoData 语义。")
    m.code("""dt_grid_clip_options o{};
o.struct_size = sizeof(o);
o.flags = DT_GRID_CLIP_CROP_TO_BOUNDS;
o.worker_count = 0;       // 自动最多 32；显式最多 64
o.tile_row_count = 64;
o.output_nodata_value = -9999.0;

dt_point3 polygon[] = {
    {500000.0, 3200000.0, 0.0},
    {500600.0, 3200050.0, 0.0},
    {500450.0, 3200500.0, 0.0},
    {499950.0, 3200400.0, 0.0},
};
dt_grid_handle clipped = nullptr;
dt_status s = dt_grid_clip_polygon(source, polygon, 4, &o, &clipped);
if (s == DT_OK) {
    dt_grid_destroy(clipped);
}""")
    m.table(
        ["模式/标志", "输出范围", "选择规则与用途"],
        [
            ("flags = 0", "保持源行列数和仿射", "多边形内保留，外部写 NoData；适合图层对齐"),
            ("CROP_TO_BOUNDS", "裁到包含入选节点的最小行列包络", "多边形内保留；减小后续显示和交换范围"),
            ("INVERT", "保持源行列数和仿射", "多边形内及边界写 NoData，外部保留；适合挖孔"),
        ],
        [2500, 3000, 3860],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    m.para("实现先用源 GRID 完整六参数逆仿射把世界坐标多边形一次变换到连续列、行空间，再对每个整数节点执行边界判断和偶奇射线判定，所以旋转、剪切和负像元高都能正确工作。紧凑裁剪会平移输出仿射原点，使输出 (0,0) 节点仍对应原 GRID 的同一世界位置。源 NoData/NaN 永远不会因裁剪变成有效值。")
    m.code("""dt_task_handle task = nullptr;
dt_grid_clip_polygon_async(source, polygon, 4, &o, &task);
// 输入多边形和 options 在创建任务时已深拷贝；可立即释放调用方数组
dt_grid_handle clipped = nullptr;
if (dt_task_get_grid_result(task, &clipped) == DT_OK) {
    dt_grid_destroy(clipped); // 调用方拥有结果句柄
}
dt_task_destroy(task);""")
    m.callout("输入、取消与原子性", "options 固定为 64 字节；至少 3 个非退化顶点，相邻点不得重复，CROP_TO_BOUNDS 与 INVERT 不可同时设置。异步创建时深拷贝多边形；工作线程按行块写互不重叠区域，只有协调线程报告进度。失败、无入选节点或取消均不发布半成品，源 GRID 保持不变。", "blue")
    m.callout("节点级边界", "本接口裁剪的是 GRID 节点值，不会沿多边形切开栅格单元；紧凑范围也按入选节点包络确定。若工程流程要求像元覆盖率或单元解析边界，应在后续矢量/栅格叠加模块中处理。", "gold")
    m.callout("本机基准", "4096×4096、8 顶点多边形保持范围掩膜：单线程 0.378 s，自动并行 0.0755 s，约 5.01×；串并行校验和误差为 0。收益受多边形复杂度、CPU、内存带宽和块高影响，不是固定承诺。", "gold")

    m.h2("4.14 GRID 窗口概览与 LOD 读取")
    m.para("dt_grid_read_overview() 从任意源窗口直接生成较小的概览矩阵，避免 GUI 或集成系统先读取整幅超大 GRID。输出数组由调用方分配，接口同步填充，不创建第二个 GRID 句柄；因此热路径只有小型输出和每输出行统计的临时开销。source_width/source_height 为 0 时表示从起点延伸到源 GRID 剩余范围。")
    m.code("""dt_grid_overview_options o{};
o.struct_size = sizeof(o);
o.method = DT_GRID_OVERVIEW_AVERAGE;
o.worker_count = 0;       // 自动最多 32；显式最多 64
o.tile_row_count = 16;

std::vector<double> preview(512 * 512);
dt_grid_overview_result r{};
r.struct_size = sizeof(r);
dt_status s = dt_grid_read_overview(
    grid, &o, 512, 512, preview.data(), 512, &r);""")
    m.table(
        ["方法", "每个输出值", "建议用途"],
        [
            ("AVERAGE（0/1）", "整数分箱内所有有效源节点的平均值", "连续高程与坡度/阴影概览；GUI 默认"),
            ("NEAREST", "取分箱中心对应的最近源节点；允许上采样", "坡向角、分类或离散值，避免角度环绕平均"),
            ("MINIMUM", "整数分箱内有效源节点最小值", "保留下包络或极低点"),
            ("MAXIMUM", "整数分箱内有效源节点最大值", "保留上包络或极高点"),
        ],
        [2100, 3950, 3210],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    m.para("聚合方法使用 floor(i×源尺寸/输出尺寸) 的整数边界，每个源节点恰好进入一个输出分箱，因而拒绝上采样；最近邻按分箱中心取样并允许上采样。默认忽略分箱内 NoData/NaN，若没有有效值则输出源 NoData（源未定义时为 NaN）；DT_GRID_OVERVIEW_STRICT_NODATA 使任一无效源节点将整个分箱标为 NoData。")
    m.para("dt_grid_overview_options 与 dt_grid_overview_result 均固定为 64 字节。平均、最小和最大方法的结果统计覆盖完整源窗口，并设置 DT_GRID_OVERVIEW_EXACT_SOURCE_STATISTICS；最近邻的统计只覆盖实际输出样本。output_result 可省略，row_stride 以 double 数量计，便于直接写入带行填充的图像或矩阵。")
    m.callout("并行与所有权", "工作线程动态领取输出行块，只写互不重叠的调用方行和统计槽；协调线程按输出行顺序归并 long double 和，保证串并行结果确定。单轴最多 1,048,576，输出总数最多 10 亿。接口仅在同步调用期间借用 grid、options 和输出数组，返回后即可安全复用或释放调用方缓存。", "blue")
    m.callout("本机基准", "8192×4096 源 GRID 缩至 512×512 平均概览：单线程 0.0427 s，自动并行 0.00963 s，约 4.43×；串并行输出与统计误差为 0。收益受 CPU、内存带宽、源窗口和块高影响，不是固定承诺。", "gold")

    m.h2("4.15 世界视口裁剪与自适应 LOD")
    m.para("dt_grid_get_view_window() 把轴对齐的世界 XY 视口转换为与 GRID 覆盖相交的最小源节点窗口。它先使用完整六参数仿射的逆变换把视口四角映射到连续列行坐标，再把所得四边形裁剪到 [-0.5,width-0.5]×[-0.5,height-0.5] 的节点覆盖域；因此旋转、剪切、负 pixel_height 以及只覆盖边缘半个像元的视口都能得到一致结果。")
    m.code("""dt_grid_view_options v{};
v.struct_size = sizeof(v);
v.world_bounds = {xmin, ymin, xmax, ymax};
v.padding_nodes = 2;       // GUI 用于避免边缘采样缝隙

dt_grid_window w{};
w.struct_size = sizeof(w);
dt_status s = dt_grid_get_view_window(grid, &v, &w);
if (s == DT_OK) {
    dt_grid_overview_options o{};
    o.struct_size = sizeof(o);
    o.method = DT_GRID_OVERVIEW_AVERAGE;
    o.source_column = w.column;
    o.source_row = w.row;
    o.source_width = w.width;
    o.source_height = w.height;
    // 继续调用 dt_grid_read_overview() 生成屏幕大小的局部 LOD
}""")
    m.para("dt_grid_view_options 与 dt_grid_window 均固定为 64 字节。padding_nodes 以源节点计并在裁剪后扩展窗口；结果设置 DT_GRID_VIEW_WINDOW_CLIPPED 表示原视口有部分超出 GRID 覆盖范围。完全无交集返回 DT_E_NOT_FOUND；未知标志、非有限或反向范围、退化仿射及过大 padding 返回 DT_E_INVALID_ARGUMENT。")
    m.callout("热路径特性", "该调用不读取高程、不分配结果数组，也不构造新 GRID；返回的 column/row/width/height 可原样复制到 dt_grid_overview_options。GUI 在滚轮缩放、拉框放大、全图适屏及平移结束后刷新，平移过程中复用旧缓存，减少交互抖动。完整源统计维持稳定色带，所有分析、等高线与导出仍读取完整 GRID。", "blue")
    m.callout("本机基准", "8192×4096 GRID、1024×768 局部视口、512×512 输出：视口窗口查询约 15 μs；全幅概览 0.0128 s，局部概览 0.00264 s，约 4.85×。局部窗口约占源节点 2.34%；收益随当前视口、硬件和内存带宽变化。", "gold")

    m.h1("5 原 12 个接口兼容层")
    m.para("include/dt_legacy.hpp 提供原需求中的 12 个 C++ 接口。它们使用 DLL 内部的全局默认上下文，适合保持既有调用代码不变。新系统优先使用 dt_api.h，以获得多上下文、明确状态码和结果句柄。")
    legacy_rows = [
        ("1", "dt_init_dll", "初始化默认上下文；与 dt_free_dll 配对。"),
        ("2", "dt_free_dll", "释放默认上下文及其资源。"),
        ("3", "dt_insert_a_point_with_draw", "插点并返回受影响面、边界和新增边。"),
        ("4", "dt_insert_a_point", "仅插入 XYZ 点。"),
        ("5", "dt_delete_a_point_with_draw", "删除查询位置 XY 最近顶点并返回影响数据。"),
        ("6", "dt_delete_a_point", "仅删除查询位置 XY 最近顶点。"),
        ("7", "dt_clear", "清空默认三角网。"),
        ("8", "dt_save_triangulation", "保存 DTIN 二进制三角网。"),
        ("9", "dt_load_triangulation", "加载 DTIN 并返回 XY 范围。"),
        ("10", "dt_view_to_range", "返回与指定 XY 矩形相交的三角形。"),
        ("11", "dt_get_a_point_nearest_point", "返回查询位置 XY 最近顶点。"),
        ("12", "dt_get_a_triangle_covers_point", "返回平面投影严格包含查询点的三角形。"),
    ]
    m.table(["序号", "函数", "行为"], legacy_rows, [850, 3650, 4860],
            [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT])
    m.h2("5.1 pEffect 数组布局")
    m.para("带绘制结果的接口输出 f、h、e 和 pEffect。pEffect 中依次放置 f 个受影响三角形、h 条边界边、e 条新增边；每个空间点由连续三个 double 表示。")
    m.code("""对象数量 = f * 3 + (h + e) * 2
double 数量 = 3 * [f * 3 + (h + e) * 2]

区段 1: f * 3 个 XYZ 点  -> 受影响三角形
区段 2: h * 2 个 XYZ 点  -> 影响区边界边
区段 3: e * 2 个 XYZ 点  -> 新增边""")
    m.callout("所有权", "旧接口返回的 double* 由 DLL 管理，调用方不得 delete/free。它只在同一线程下一次 Legacy 接口调用前有效，需要长期保存时必须立即复制。", "red")
    m.h2("5.2 Legacy 最小调用流程")
    m.code("""#include "dt_legacy.hpp"

dt_init_dll();
dt_insert_a_point(0.0, 0.0, 10.0);
dt_insert_a_point(100.0, 0.0, 12.0);
dt_insert_a_point(0.0, 100.0, 20.0);

double x, y, z;
bool ok = dt_get_a_point_nearest_point(
    x, y, z, 10.0, 15.0, 0.0);

dt_free_dll();""")

    m.h1("6 文本与二进制格式")
    m.h2("6.1 XYZ 散点文本")
    m.para("每个有效行恰好包含三个有限浮点数 x y z。支持空格、制表符、逗号、分号、空行、整行 # 注释和第三个坐标后的行尾注释。")
    m.code("""# x, y, z
500000.0, 3200000.0, 103.25
500010.0 3200000.0 104.10
500010.0;3200010.0;106.80 # ridge""")
    m.h2("6.2 DTMESH 1")
    m.code("""DTMESH 1
VERTICES 4
1 0 0 10
2 1 0 20
3 1 1 30
4 0 1 40
TRIANGLES 2
1 2 3
1 3 4""")
    for text in (
        "VERTICES 每行是 id x y z；ID 必须为非零 uint64_t。",
        "TRIANGLES 每行是三个顶点 ID。",
        "加载时从顶点重建 Delaunay，并检查面数、引用、重复面与拓扑。",
        "不匹配时返回 DT_E_CORRUPTED_DATA，原网保持不变。",
    ):
        m.bullet(text)
    m.h2("6.3 DTIN v1")
    m.para("DTIN v1 保存顶点 ID 和 XYZ，加载时重建有效 Delaunay 网。对于共圆点集，重建后可能选择另一条同样合法的对角线，因此 DTIN v1 不承诺逐面拓扑完全一致。需要文本可读和逐面校验时使用 DTMESH。")
    m.h2("6.4 DGRID 规则高程节点文本")
    m.para("DGRID 1 保存 SIZE、FLAGS、六参数 GEOTRANSFORM、NODATA 和按行优先排列的 VALUES。该格式用于测试、研究和简单交换；当前变换描述节点位置，不等同于 GDAL 像元左上角语义。")
    m.h2("6.5 DCONTOUR 等高线文本")
    m.para("DCONTOUR 1 由 LINES 计数和若干 LINE 记录组成。每条 LINE 保存等高值、闭合标志、点数及 XYZ 顶点。逐线视图中的点数组只在等高线句柄释放前有效。")
    m.h2("6.6 GeoTIFF、COG 与 GeoPackage")
    m.para("启用 DT_WITH_GDAL 后，dt_grid_load/save_gdal_raster 读写 GDAL 单波段高程栅格；dt_contours_load/save_gdal_vector 读写等高线矢量层。普通 GeoTIFF 逐行写入；COG 按驱动要求通过临时数据集 CreateCopy；GeoPackage 输出 LineStringZ 和 elevation/closed 字段。")
    m.h2("6.7 DCDT 约束网文本")
    m.para("DCDT 1 保存 CRS、基础 XYZ 点和约束记录。每条约束包含稳定 ID、类型、闭合标志、点数和 XYZ；加载时完整重建 CDT 并重新标记域，失败不会覆盖原对象。")

    m.h1("7 错误、内存与线程")
    m.h2("7.1 状态码")
    status_rows = [
        ("DT_OK", "成功"),
        ("DT_E_INVALID_ARGUMENT", "空指针、无效坐标、无效矩形或参数组合"),
        ("DT_E_NOT_INITIALIZED", "上下文或 Legacy 默认环境尚未初始化"),
        ("DT_E_DUPLICATE_XY", "输入或插入点与现有顶点具有相同 XY"),
        ("DT_E_EMPTY", "在空网格上执行需要数据的操作"),
        ("DT_E_NOT_FOUND", "目标顶点或对象不存在"),
        ("DT_E_IO", "文件打开、读取或写入失败"),
        ("DT_E_OUT_OF_MEMORY", "内存分配失败"),
        ("DT_E_CORRUPTED_DATA", "文件格式、计数或拓扑校验失败"),
        ("DT_E_STALE_QUERY", "结果与当前 generation 不一致"),
        ("DT_E_INTERNAL", "未分类内部异常"),
        ("DT_E_UNSUPPORTED", "当前后端不支持所请求语义，例如未允许的 NoData 空洞"),
        ("DT_E_CANCELLED", "异步地形任务收到协作取消请求"),
        ("DT_E_LIMIT_EXCEEDED", "GRID、等高层或文本对象数量超过安全限制"),
    ]
    m.table(["状态", "说明"], status_rows, [3200, 6160],
            [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT])
    m.h2("7.2 错误文本")
    m.code("""char message[1024]{};
size_t required = 0;
dt_get_last_error(message, sizeof(message), &required);""")
    m.para("最后错误文本按线程保存。接口失败后应在同一线程尽快读取；后续 DLL 调用可能覆盖它。")
    m.h2("7.3 结果所有权")
    m.table(
        ["对象", "创建者", "释放方式"],
        [
            ("dt_handle", "dt_create", "dt_destroy"),
            ("dt_edit_result", "插入/删除接口", "dt_release_edit_result"),
            ("dt_query_result", "dt_query_triangles", "dt_release_query_result"),
            ("dt_grid_handle", "GRID 创建/加载/转换", "dt_grid_destroy"),
            ("dt_contour_handle", "等高线生成/加载", "dt_contours_destroy"),
            ("dt_task_handle", "异步转换启动", "dt_task_destroy"),
            ("dt_cdt_handle", "dt_cdt_create", "dt_cdt_destroy"),
            ("dt_cdt_query_result", "dt_cdt_query_triangles", "dt_cdt_release_query_result"),
            ("Legacy double*", "兼容层", "不得释放；立即复制"),
        ],
        [2500, 3100, 3760],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    m.h2("7.4 并发")
    m.para("每个上下文有读写锁。但当前 MinGW + CGAL 6.0.1 构建由于 TLS ABI 兼容性问题定义 CGAL_HAS_NO_THREADS，并用全局递归互斥锁串行保护 CGAL 调用。因此接口线程安全不等于能够并行扩展吞吐。Legacy 接口还共享一个默认上下文。")
    m.callout("建议", "同一网格的业务写操作在应用层串行调度；查询线程应避免持有结果对象过久。需要跨网格并行时，优先评估 MSVC/更新 CGAL 工具链并重新进行并发验证。", "blue")

    m.h1("8 大数据性能与容量规划")
    m.table(
        ["点数", "有限三角形", "建网时间", "峰值工作集", "完整校验"],
        [
            ("100,000", "199,973", "0.21 s", "未单独记录", "0.02 s"),
            ("1,000,000", "1,999,962", "3.47 s", "460.5 MB", "0.36 s"),
            ("10,000,000", "19,999,951", "50.44 s", "4,553.2 MB", "3.39 s"),
        ],
        [1700, 1900, 1700, 2200, 1860],
        [WD_ALIGN_PARAGRAPH.CENTER] * 5,
    )
    m.para("数据来自 Windows x64 Release/MinGW、随机均匀 XY 点集，用于说明数量级，不是对任意硬件和坐标分布的固定承诺。")
    m.table(
        ["CDT 场景", "逐条调用", "批量事务", "本次加速"],
        [("100,000 点 + 12 条断裂线", "3.259 s", "0.528 s", "6.17×")],
        [3500, 1900, 1900, 2060],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
         WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER],
    )
    m.h2("8.1 使用建议")
    for text in (
        "百万级以上优先批量构建，提前验证可用内存和临时峰值。",
        "使用局部或投影坐标；避免 NaN、Inf 和重复 XY。",
        "按视口调用矩形查询，不要为显示一次性导出全部千万级三角形。",
        "保存交换时，DTIN 更紧凑；DTMESH 更可读但体积更大。",
        "全量 dt_validate() 适合导入验收、调试和离线检查，不必每帧调用。",
    ):
        m.bullet(text)

    m.h1("9 测试与验收")
    m.h2("9.1 自动化测试")
    m.code("""cmake --build build --config Release --parallel 4
ctest --test-dir build --output-on-failure""")
    m.para("测试覆盖生命周期、普通 TIN 动态编辑、DTIN/DTMESH/XYZ、TIN/GRID/等高线、等高线反向重建与高程冲突回滚、异步任务，以及 CDT 空网清理、外边界、孔洞、断裂线、交叉失败原子性、TIN→CDT、域内采样、CDT→GRID/等高线、共享顶点保护、批量新增/更新/删除、整批失败回滚和 DCDT 往返。局部坡面测试用解析平面验证 TIN/CDT 的 Z、世界梯度、坡度、坡向和法向；全幅测试覆盖坡度、坡向、阴影、CRS、旋转仿射、NoData、串并行一致和取消。土方测试锁定 64/112 字节 ABI、零线穿越解析体积、旋转面积、NoData、差值 GRID、异步结果和取消。v0.24 重采样测试另锁定 64 字节 ABI、旋转/剪切解析平面、最近邻与双线性、严格/归一化 NoData、串并行一致、异步结果、CRS 拒绝及大 GRID 取消。v0.25 多边形裁剪测试锁定 64 字节 ABI、旋转/剪切仿射、边界内含、保持/紧凑/反向三种模式、NoData、串并行一致、异步结果、非法选项、范围外与大 GRID 取消。v0.26 概览测试锁定 64/64 字节 ABI、平均/最近邻/最小/最大、严格 NoData、源子窗口、上采样规则、带填充步长以及串并行确定性。v0.27 视口窗口测试锁定 64/64 字节 ABI、负像素高、旋转/剪切、边界裁剪、padding、无交集与非法参数，并验证结果窗口可直接组合概览读取。剖面、多边形量测和 GDAL 构建另有数学与真实文件往返测试。")
    m.h2("9.2 集成验收清单")
    for text in (
        "DLL、导入库、头文件和运行库均为相同架构。",
        "三点、四点、重复 XY、共线点、空文件和非法文件均有预期结果。",
        "所有结果句柄在成功和失败路径上都得到释放。",
        "中文路径通过 UTF-8 传入新 API。",
        "百万级样本的建网时间、内存和局部查询满足目标系统预算。",
        "保存后重新加载，顶点数、范围、高程和查询结果通过核对。",
        "使用解析平面核对三种坡面接口的梯度、坡度、下坡坡向与单位法向约定。",
    ):
        m.bullet(text)

    m.h1("10 常见问题")
    faq = [
        ("DLL 无法加载", "检查 dterrain.dll、libgmp-10.dll、libwinpthread-1.dll 是否同目录，并确认 x64/x86 一致。"),
        ("链接不到符号", "包含 dt_api.h，并链接 libdterrain.dll.a 或 dterrain.lib；不要混用 MSVC 与 MinGW 导入库。"),
        ("导入返回重复 XY", "同一 XY 只能保留一个高程；先清洗数据或明确采用何种聚合规则。"),
        ("文本三角网打不开", "确认 DTMESH 1 头、计数、ID 引用和三角形集合均与 Delaunay 拓扑一致。"),
        ("全图查询内存大", "缩小 dt_bounds2 范围，按视口或瓦片分批查询。"),
        ("Z 改变后面没有变化", "这是设计行为：Delaunay 只使用 XY，Z 更新不改变拓扑。"),
        ("插入点失败", "检查有限坐标、重复 XY、初始化状态和 dt_get_last_error()。"),
        ("GDAL 接口不支持", "重新配置 DT_WITH_GDAL=ON，并确认 GDAL::GDAL 可被 CMake 找到。"),
        ("GDAL 版 DLL 无法加载", "将构建目录复制出的 GDAL/PROJ/SQLite/GeoTIFF 等依赖 DLL 与 dterrain.dll 放在同目录。"),
    ]
    m.table(["现象", "处理"], faq, [2600, 6760],
            [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT])

    m.h1("11 许可证")
    m.para("当前后端使用 CGAL 2D Triangulations 包。该包采用 GPL 或商业许可双重模式；本工程定位为研究和演示用途。对外分发、闭源集成或用途改变前，应再次核查 CGAL 及依赖许可。")

    path = OUTPUT / "dterrain_DLL开发使用手册.docx"
    m.save(path)
    return path


def build_gui_manual():
    m = Manual(
        "dterrain GUI 演示程序\n操作手册与入门教程",
        "GUI 操作入门教程",
        "从 XYZ 散点导入到动态编辑、大 GRID 自动 LOD、任意多边形 GRID 裁剪、错位设计 GRID 显式对齐、现状/设计双表面土方、全幅地形专题图、局部坡面、剖面与多边形量测、TIN/GRID/等高线转换、约束 Delaunay、三维漫游与数据保存",
        "首次使用者、演示人员、算法验证与测绘研究人员",
    )
    m.cover("QUICK START GUIDE")
    add_document_control(
        m,
        "便携运行、五分钟上手、控件与菜单、XYZ 导入、TIN/GRID/等高线/CDT 图层、视口自适应大 GRID LOD、任意多边形 GRID 裁剪/掩膜、错位 GRID 最近邻/双线性对齐、现状/设计双 GRID 土方、全幅坡度/坡向/阴影专题图、局部坡面、任意剖面、多边形面积/土方、2D/3D 浏览、查询编辑和数据保存。",
        [
            "第一次体验：直接完成第 2 章的五分钟教程。",
            "已有 XYZ 数据：重点阅读第 5 章。",
            "三维地形展示：重点阅读第 4 章。",
            "大 GRID 视口自适应 LOD：重点阅读第 4.4 节。",
            "任意剖面分析：重点阅读第 6.3 节。",
            "面积与土方量测：重点阅读第 6.4 节。",
            "任意区域 GRID 裁剪：重点阅读第 6.5 节。",
            "现状/设计双表面土方：重点阅读第 6.6 节。",
            "错位设计 GRID 对齐：重点阅读第 6.6 节。",
            "坡度、坡向与法向：重点阅读第 6.7 节。",
            "全幅地形专题图：重点阅读第 6.8 节。",
            "TIN/GRID/等高线转换与交换：重点阅读第 4.3 节和第 8 章。",
            "演示动态编辑：重点阅读第 7 章。",
            "百万级数据展示：重点阅读第 9 章。",
        ],
    )

    m.h1("1 程序概览")
    m.para("dterrain_demo.exe 是一个原生 Win32/GDI 演示程序，用于直观验证 dterrain.dll 的批量建网、范围查询、最近点查询、大 GRID 视口自适应 LOD、任意多边形 GRID 裁剪/掩膜、错位设计 GRID 显式对齐、现状/设计双 GRID 挖填方、全幅坡度/坡向/阴影专题图、单点坡度/坡向与法向、任意地形剖面、多边形面积/土方量测、动态插入删除、编辑影响显示、文件交换和三维地形漫游。它不依赖 Qt 等 GUI 框架。")
    m.callout("0.27 功能边界", "当前程序根据二维世界视口裁剪 GRID 源窗口，再把该局部窗口生成最多 512×512 的显示 LOD；缩放、拉框、全图适屏及平移结束都会自动刷新。高程、坡度、阴影和差值用平均，坡向用最近邻；状态栏报告“源窗口→预览”尺寸。完整 GRID 仍用于等高线、分析、转换与导出。程序支持旋转/剪切仿射但不会隐式重投影；磁盘流式瓦片、持久化金字塔、单元解析边界裁剪和 GPU LOD 尚未接入。", "gold")
    m.h2("1.1 运行文件")
    m.table(
        ["文件", "作用"],
        [
            ("dterrain_demo.exe", "GUI 主程序。"),
            ("dterrain.dll", "动态 Delaunay TIN 功能。"),
            ("libgmp-10.dll", "CGAL/GMP 数值运行库。"),
            ("libwinpthread-1.dll", "MinGW 线程运行库。"),
            ("sample_data/sample_points.xyz", "随包提供的入门散点样例。"),
            ("sample_data/sample_grid.dgrid", "随包提供的规则 GRID 文本样例。"),
            ("sample_data/sample_contours.dcontour", "随包提供的等高线文本样例。"),
            ("sample_data/sample_constraints.dcdt", "含外边界、孔洞和断裂线的 CDT 样例。"),
        ],
        [3700, 5660],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    m.callout("启动", "打开 dist/bin/dterrain_demo.exe。四个运行文件应位于同一 bin 目录；无需安装程序。", "blue")
    m.h2("1.2 启动后的默认状态")
    m.para("程序启动时会自动生成约 10 万个模拟地形点并构网。画布显示按高程分级着色的三角形线框，状态栏给出当前模式、顶点数、三角形数、当前窗口查询数量和耗时。")

    m.h1("2 五分钟上手")
    m.callout("目标", "导入示例 XYZ，浏览和编辑 TIN，生成并按多边形裁剪 GRID，体验等高线与三维漫游，最后保存数据。", "gray")
    steps = [
        "启动 dist/bin/dterrain_demo.exe，等待默认 10 万点网格显示完成。",
        "单击“导入XYZ”，选择 dist/sample_data/sample_points.xyz。导入成功后程序立即自动构网并全图适屏。",
        "滚动鼠标滚轮放大或缩小；按住右键拖动进行平移。",
        "单击“框选放大”，在画布中按住左键拖出矩形，松开后选区自动适配窗口。",
        "单击“切换3D”，左键拖动环视、滚轮拉近，再用 WASD 漫游；单击“高程×1.0”观察垂直夸张。",
        "单击“切换2D”返回二维编辑视图。",
        "执行“地形转换→从当前 TIN 创建约束网”，得到没有约束的 CDT。",
        "选择“约束编辑→批量添加 12 条示例断裂线”，观察一次事务完成后的彩色平行线和状态栏耗时。",
        "打开“地形转换”菜单，依次执行“TIN → GRID”和“从 GRID 生成等高线”，观察三类图层叠加。",
        "继续执行“等高线 → TIN”和“等高线 → GRID”，观察源等高线保留并与近似重建结果叠加。",
        "打开“图层”菜单，分别隐藏和显示 TIN、GRID、等高线，再用“全图”恢复联合范围。",
        "若当前为 GDAL 构建，执行‘数据交换→导出 GeoTIFF（DEFLATE）’和‘导出 GeoPackage 等高线’，再重新导入检查图层叠加。",
        "打开“数据交换→打开约束网 DCDT”，选择 sample_constraints.dcdt，观察孔洞和断裂线。",
        "打开“约束编辑→绘制断裂线”，在画布逐点单击，按 Enter 完成；用 Backspace 撤点、Esc 取消。",
        "选择“移动约束顶点（两次单击）”，第一次单击白色约束顶点使其变黄，第二次单击目标位置；观察红色旧面、黄色边界和绿色新增面/边。",
        "选择“删除约束顶点（单击）”并拾取白色折点；共享顶点确认后只从当前约束脱离。",
        "执行“地形转换→约束网 → GRID”和“从约束网生成等高线”，确认孔洞区域保持空白。",
        "单击“查询模式”，在网格内单击；观察白色最近顶点和洋红色覆盖三角形。",
        "选择‘分析→坡度/坡向分析（单击）’，在坡面上单击；观察青色支撑单元、黄色查询点、下坡箭头和右上结果面板。",
        "先用‘分析→设置专题分析与性能参数’检查 z-factor、光照角、线程数和块高，再生成全幅坡度、坡向和阴影地形；观察进度、固定色带和右上图例，并用‘恢复显示高程 GRID’返回高程着色。",
        "选择‘分析→任意剖面（两次单击）’，在画布选取 A、B；观察青色定位线、底部曲线和统计，再导出 CSV。",
        "选择‘分析→面积/土方量测（逐点）’，输入基准高程，逐点圈定区域并按 Enter；核对顶部结果面板，再修改基准高程和导出 CSV。",
        "先保存原 GRID，再选择‘分析→按量测多边形裁剪当前 GRID（紧凑适屏）’；观察范围缩小和自动适屏。当前无撤销命令，需要恢复时重新打开原 GRID。",
        "选择‘分析→按高程偏移创建设计 GRID’，输入 2；可直接执行‘计算现状面—设计面挖填方’，观察差值图和结果，最后导出土方 CSV。若另行打开同 CRS 但错位的设计 GRID，先执行‘双线性对齐设计面到现状 GRID’。",
        "单击“插入模式”，在网格中单击；观察红色旧面、黄色边界和绿色新增面/边。",
        "单击“删除模式”，在目标附近单击；程序删除 XY 最近顶点并显示局部变化。",
        "单击“保存网格”，保存为 terrain.dtmesh；然后单击“清空”和“打开网格”验证加载。",
    ]
    for step in steps:
        m.number(step)
    m.callout("完成标志", "状态栏可显示 TGCD；画布能叠加普通 TIN、GRID、等高线与约束网，并能重新打开 .dtmesh、.dgrid、.dcontour 或 .dcdt。", "blue")

    m.h1("3 界面组成")
    m.h2("3.1 顶部工具栏")
    control_rows = [
        ("模拟10万", "生成约 100,000 个起伏地形点并自动构网。"),
        ("模拟100万", "生成约 1,000,000 个点，用于大数据交互演示。"),
        ("清空", "删除当前网格全部数据。"),
        ("查询模式", "单击显示最近顶点和查询点所在三角形。"),
        ("插入模式", "单击插入一个按演示地形函数计算 Z 的新点。"),
        ("删除模式", "删除单击位置 XY 距离最近的顶点。"),
        ("框选放大", "左键拖出矩形，松开后按画布比例放大适屏。"),
        ("全图", "恢复到完整数据范围并适屏。"),
        ("导入XYZ", "导入 XYZ/TXT/CSV 散点并立即构网。"),
        ("保存网格", "保存 DTMESH 文本或 DTIN 二进制网格。"),
        ("打开网格", "打开并校验 DTMESH/TXT，或加载 DTIN。"),
        ("切换3D/切换2D", "在二维编辑视图和透视三维 TIN 视图之间切换。"),
        ("高程×n", "循环选择 0.5、1.0、1.5、2、3、5、8 倍垂直夸张。"),
    ]
    m.table(["控件", "功能"], control_rows, [2600, 6760],
            [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT])
    m.h2("3.2 主画布与状态栏")
    m.para("主画布绘制当前视口范围内的地形图层。底部状态栏从左到右显示：2D/3D 视图、当前模式、顶点数、有限三角形数、当前窗口查询结果数、查询耗时或垂直夸张、T/G/C/D 可见图层标记，以及最近一次操作说明。")
    m.h2("3.3 菜单栏")
    m.table(
        ["菜单", "主要命令"],
        [
            ("地形转换", "TIN/GRID/等高线双向派生，CDT 域内转换，并自动生成等高线。"),
            ("数据交换", "交换 DGRID、DCONTOUR、DCDT 文本；GDAL 构建还支持 GeoTIFF/COG 与 GeoPackage。"),
            ("约束编辑", "绘制或批量添加断裂线；移动/安全删除顶点，或删除整条约束。"),
            ("分析", "生成全幅坡度/坡向/阴影专题图并导出；单击分析局部坡面；生成任意 A—B 剖面；逐点圈定多边形并估算面积/土方。"),
            ("图层", "分别显示或隐藏 TIN、GRID、等高线和约束 Delaunay。"),
        ],
        [2600, 6760],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
    )

    m.h1("4 浏览与视图控制")
    m.table(
        ["操作", "输入", "结果"],
        [
            ("滚轮缩放", "在画布中滚动滚轮", "以鼠标位置为中心连续缩放。"),
            ("平移", "按住右键或中键拖动", "移动当前观察范围。"),
            ("框选放大", "选择模式后左键拖框", "选区放大并保持 XY 比例。"),
            ("取消框选", "拖动过程中按 Esc", "取消选择，不改变视图。"),
            ("恢复全图", "单击“全图”", "数据完整范围重新适屏。"),
        ],
        [2200, 3300, 3860],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    m.h2("4.1 框选放大细节")
    for text in (
        "选择“框选放大”后，鼠标左键用于选择，不执行查询或编辑。",
        "拖动时显示青色虚线框；可以从任意方向拖动。",
        "松开后程序按画布纵横比扩展选区，避免地形图形被拉伸。",
        "宽或高小于 8 像素的误拖会被忽略。",
        "放大后仍保持框选模式，可连续选择更小区域；单击其他模式可退出。",
    ):
        m.bullet(text)

    m.h2("4.2 三维查看与漫游")
    m.para("单击“切换3D”后，完整 TIN 以透视方式显示；红、绿、蓝短轴分别代表 X、Y、Z。三维只负责查看，选择查询、插入、删除或框选放大时会自动返回二维，防止把透视屏幕坐标误当作地形 XY。")
    m.table(
        ["输入", "三维操作"],
        [
            ("左键拖动", "绕观察目标环视。"),
            ("右键/中键拖动", "沿相机屏幕平面平移。"),
            ("滚轮", "拉近或拉远相机。"),
            ("W/S 或 ↑/↓", "沿当前水平方向前进或后退。"),
            ("A/D 或 ←/→", "向左或向右漫游。"),
            ("Q/E", "提升或降低观察目标。"),
            ("+/-", "增加或降低垂直夸张。"),
            ("Home / 全图", "恢复三维适屏相机。"),
        ],
        [3200, 6160],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
    )
    m.callout("大坐标稳定性", "三维模型先减去数据中心并按 XY 范围归一化，再使用双精度相机与投影计算，可直接查看局部坐标或大数值投影坐标。", "blue")

    m.h2("4.3 TIN、GRID、等高线与 CDT 图层")
    m.para("TIN→GRID 自动按 TIN 范围建立最长边 401 节点规则网；TIN/GRID/CDT→等高线自动选择易读等高距。等高线→TIN 使用折点与线高程重建普通 Delaunay，等高线→GRID 经临时 TIN 插值；两者都保留源等高线。GRID→TIN 仍可能跨 NoData；CDT→GRID 把外边界以外和孔洞内部写为 NoData，CDT 等高线不会跨越无效域。‘从当前 TIN 创建约束网’会复制顶点与 CRS，并清除原 CDT 约束。")
    m.table(
        ["图层", "二维显示", "数据变化规则"],
        [
            ("TIN", "高程分级网线和编辑效果", "重新建网或动态编辑会使派生 GRID/等高线失效。"),
            ("GRID", "连续高程着色和青色边框", "GRID→TIN 保留源 GRID。"),
            ("等高线", "黄色普通线与浅色加粗抽样线", "可反向近似重建 TIN/GRID；不恢复丢失极值。"),
            ("CDT", "紫色域内网、青色外边界、洋红孔洞、橙色断裂线", "可打开、保存、交互编辑并派生 GRID/等高线。"),
        ],
        [1600, 3000, 4760],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    m.callout("全图与显隐", "“全图”按所有可见二维图层的联合范围适屏。状态栏 T/G/C/D 分别表示普通 TIN、GRID、等高线和约束 Delaunay。", "blue")

    m.h2("4.4 大 GRID 的视口自适应 LOD 预览")
    m.para("打开 DGRID/GeoTIFF、执行 TIN→GRID 或生成地形专题图后，程序先取得完整 GRID 的稳定高程范围；随后把当前二维世界视口映射为最小源节点窗口，并仅把这个局部窗口生成最长边不超过 512 节点的内存概览。滚轮缩放、拉框放大、全图适屏和拖动平移结束时都会自动刷新；拖动过程中复用旧图，松开后再细化，保持交互连续。状态栏会给出类似“局部LOD 1024×768→512×384”的源窗口和预览尺寸。")
    m.table(
        ["当前显示", "概览方法", "原因"],
        [
            ("高程 GRID", "平均", "连续高程在缩小时保持整体形态，并使用完整源窗口精确极值着色"),
            ("坡度 / 阴影 / 差值", "平均", "适合连续专题值，减少缩小时的锯齿与闪烁"),
            ("坡向", "最近邻", "坡向是环形角度，普通平均会破坏 0°/360° 附近值"),
        ],
        [2200, 1850, 5210],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
    )
    m.callout("仿射与边缘", "视口裁剪使用完整六参数仿射逆变换，支持旋转、剪切和负像素高。程序在局部窗口四周增加 2 个源节点，减少缩放与平移后的边缘缝隙；GRID 的青色轮廓始终按完整范围绘制。", "gold")
    m.callout("数据完整性", "LOD 只服务二维着色缓存；色带使用完整源统计保持缩放前后稳定。保存 DGRID/GeoTIFF、GRID→TIN、裁剪、分析和从 GRID 生成等高线仍使用完整 GRID。即使源 GRID 有数千万节点，也不会因为屏幕预览而丢值。", "blue")
    m.callout("本机参考", "8192×4096 GRID 在 1024×768 局部视口生成 512×512 概览约 0.00264 s；同尺寸输出从全幅生成约 0.0128 s，约快 4.85×。视口裁剪查询本身约 15 μs。实际速度由缩放级别、CPU 和内存带宽决定。", "gray")

    m.h1("5 导入 XYZ 散点并自动构网")
    m.h2("5.1 支持格式")
    m.para("每个有效行包含 x、y、z 三个有限数值。支持空格、制表符、逗号和分号分隔；支持空行、整行 # 注释和行尾注释。文件扩展名可为 .xyz、.txt 或 .csv。")
    m.code("""# x, y, z
500000.0, 3200000.0, 103.25
500010.0 3200000.0 104.10
500010.0;3200010.0;106.80 # ridge""")
    m.h2("5.2 导入步骤")
    for step in (
        "单击“导入XYZ”。",
        "在文件对话框中选择散点文件。首次体验可选择 dist/sample_data/sample_points.xyz。",
        "等待鼠标恢复普通指针；大文件构网期间程序会显示等待指针。",
        "查看状态栏中的点数和耗时，确认网格自动全图适屏。",
        "使用查询模式抽查高程和覆盖三角形。",
    ):
        m.number(step)
    m.h2("5.3 导入失败的常见原因")
    m.table(
        ["提示/现象", "原因与处理"],
        [
            ("contains no points", "文件为空、全是注释或没有有效 XYZ 行。"),
            ("invalid XYZ", "某行字段不足、多余、不是数字或存在 NaN/Inf。"),
            ("duplicate XY", "两个点的 X 和 Y 相同；先清洗或聚合高程。"),
            ("cannot open", "路径、权限、文件占用或文件名转换问题。"),
            ("内存不足", "减少点数，关闭其他程序，使用 64 位 Release 版本。"),
        ],
        [3000, 6360],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    m.callout("安全性", "导入具有事务性：解析或建网失败时，原三角网不会被部分覆盖。", "blue")

    m.h1("6 查询地形对象")
    m.h2("6.1 最近顶点和覆盖三角形")
    for step in (
        "单击“查询模式”。",
        "在网格内单击目标位置。",
        "白色圆点表示 XY 平面最近顶点。",
        "洋红色粗线表示查询点所在或相交的三角形。",
        "状态栏显示查询 XY、最近顶点 ID、Z 和操作耗时。",
    ):
        m.number(step)
    m.callout("注意", "查询使用 XY 投影距离，鼠标点击没有输入 Z；状态栏显示的是最近网格顶点自身的 Z。", "gold")
    m.h2("6.2 凸包外查询")
    m.para("在三角网凸包外单击时，仍可能找到最近顶点，但没有严格覆盖查询点的有限三角形。此时不会显示洋红色覆盖面。")
    m.h2("6.3 任意地形剖面")
    for step in (
        "选择‘分析→任意剖面（两次单击）’；程序自动返回二维视图。",
        "在画布单击剖面起点 A，再单击终点 B；程序沿直线建立 401 个等距样本。",
        "观察青色 A—B 定位线和端点；底部深色面板中的黄色线为高程曲线。",
        "查看状态栏或面板中的平距、有效样本、高程范围、净高差、累计上升/下降和最大绝对坡度。",
        "选择‘分析→导出剖面 CSV’，保存距离、XYZ、有效标志和相邻坡度。",
        "完成后再次单击可开始新剖面；第一点选取后按 Esc 取消，或用‘清除剖面’清除全部结果。",
    ):
        m.number(step)
    m.table(
        ["顺序", "数据源选择", "无效区域"],
        [
            ("1", "当前可见 CDT 有效域", "域外或孔洞保留断点"),
            ("2", "当前可见普通 TIN", "凸包外保留断点"),
            ("3", "当前可见 GRID", "NoData 保留断点"),
            ("4", "再按 CDT、TIN、GRID 考虑隐藏但存在的图层", "仍不逐点混用其他表面"),
        ],
        [1200, 4200, 3960],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    m.callout("一致性规则", "整条剖面只选一次数据源，断点不会用其他图层补齐。坡度按 ΔZ/Δ平距×100% 计算，断点两侧不跨段累计。GRID 使用六参数仿射反算和邻近 1×1/2×2 小窗口读取，不复制整幅大栅格。", "blue")
    m.code("index,distance,x,y,z,valid,grade_percent\n0,0.000000,500000.000000,3200000.000000,103.250000,1,\n1,2.500000,500002.500000,3200000.000000,103.650000,1,16.000000")

    m.h2("6.4 多边形面积与土方量测")
    for step in (
        "选择‘分析→面积/土方量测（逐点）’；在弹窗输入水平基准高程。",
        "在二维画布沿简单多边形边界依次单击；至少 3 点后按 Enter 闭合并计算。",
        "草图输入时按 Backspace 撤回末点；按 Esc 或菜单命令清除整块量测。",
        "观察橙色边界、黄色顶点和斜线填充；顶部面板报告平面面积、有效覆盖、地表面积、高程和挖填方。",
        "选择‘设置量测基准高程…’修改水平面；程序立即以相同边界和固定数据源重算。",
        "选择‘导出面积/土方 CSV…’，保存摘要、积分单元数和边界顶点。",
    ):
        m.number(step)
    m.table(
        ["结果", "计算口径", "说明"],
        [
            ("平面面积/周长", "输入 XY 边界", "边界几何量"),
            ("有效平面面积/覆盖率", "有效微三角形中心", "排除 CDT 域外/孔洞、TIN 凸包外和 GRID NoData"),
            ("地表面积/平均高程", "有效微三角形积分", "平均高程按平面面积加权"),
            ("挖方/填方", "相对水平基准面积分", "现状高于基准为挖方，低于基准为填方"),
            ("净挖方", "挖方 − 填方", "负值表示净填方"),
        ],
        [2100, 3000, 4260],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    m.callout("精度边界", "平面面积和周长由边界直接计算；其他三维量使用约 20,000 个微三角形数值积分。孔洞/NoData 边缘按微三角形中心分类，结果适合研究和方案比较，不应直接作为工程结算量。整块量测只使用首点确定的一个 CDT/TIN/GRID 数据源。", "gold")
    m.code("# source,CDT\n# datum_z,100.000000\n# polygon_area,2500.000000\n# valid_plan_area,2375.000000\n# cut_volume,820.500000\n# fill_volume,136.250000\nindex,x,y,z,valid")

    m.h2("6.5 按量测多边形裁剪或掩膜 GRID")
    for step in (
        "准备当前高程 GRID；若只有 TIN/CDT，先通过‘地形转换’生成 GRID。建议先保存原 GRID，因为裁剪成功后会替换当前图层。",
        "选择‘分析→面积/土方量测（逐点）’，在二维画布逐点圈定一个简单多边形，按 Enter 闭合；本功能只复用其 XY 边界，量测基准高程不影响裁剪。",
        "需要保持原行列与空间对齐时，选择‘按量测多边形掩膜当前 GRID（保持范围）’；多边形外节点写 NoData。",
        "希望减少数据范围并自动放大时，选择‘按量测多边形裁剪当前 GRID（紧凑适屏）’；结果裁到入选节点的最小行列包络。",
        "需要在 GRID 中挖出多边形孔洞时，选择‘按量测多边形反向掩膜当前 GRID’；内部和边界写 NoData，外部保留。",
        "观察状态栏进度；计算中可按 Esc 取消。成功后程序清除量测草图以及依赖旧 GRID 的等高线、设计面、土方和专题结果，并显示新 GRID。",
        "用‘数据交换’或‘保存 GRID’导出结果，也可继续执行 GRID→TIN、生成等高线、专题分析和 3D 浏览。",
    ):
        m.number(step)
    m.table(
        ["命令", "几何范围", "典型用途"],
        [
            ("保持范围掩膜", "源 GRID 行列和仿射不变", "继续与同节点设计面或其他栅格叠加"),
            ("紧凑裁剪", "入选节点的最小行列包络", "局部展示、交换和后续计算"),
            ("反向掩膜", "源 GRID 行列和仿射不变", "挖孔、排除保护区或无效区"),
        ],
        [2700, 3300, 3360],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    m.callout("边界与坐标", "边界节点按内部处理；多边形顺、逆时针均可。程序使用 GRID 完整六参数逆仿射，旋转、剪切和负像元高同样适用。裁剪只改变节点有效性，不沿多边形切开 GRID 单元。", "blue")
    m.callout("替换与恢复", "成功结果会原子替换当前 GRID，且 GUI 暂无撤销栈；请先保存重要原始数据，需要恢复时重新打开原 GRID。取消、非法多边形或无入选节点时旧图层保持不变。", "gold")

    m.h2("6.6 现状/设计双 GRID 挖填方")
    for step in (
        "准备现状高程 GRID：可打开 DGRID，GDAL 构建也可打开 GeoTIFF；当前高程 GRID 就是现状面。",
        "准备设计面：选择‘分析→打开设计 GRID（DGRID）’或‘打开设计 GRID（GDAL）’。CRS WKT 必须与现状面一致；行列数或六参数仿射不同也可先加载。",
        "快速演示时可选择‘按高程偏移创建设计 GRID…’，输入设计面相对现状面的统一高程偏移；NoData 会原样保留。",
        "若状态栏提示节点未对齐，连续高程面选择‘双线性对齐设计面到现状 GRID’，分类值或快速预览可选择‘最近邻对齐设计面到现状 GRID’。计算期间可按 Esc 取消；旧设计面会保留。",
        "选择‘计算现状面—设计面挖填方’。计算期间状态栏显示进度；按 Esc 可安全取消，旧结果不会被半成品替换。",
        "完成后查看状态栏中的挖方、填方、净方、覆盖率和耗时。画布自动显示差值 GRID：红色为现状高于设计，蓝色为现状低于设计，白色接近零。",
        "选择‘导出双表面土方 CSV…’保存汇总统计；需要更换设计时重新打开或创建，选择‘清除设计面及土方结果’可释放设计与差值图层。",
    ):
        m.number(step)
    m.table(
        ["项目", "双 GRID 土方", "多边形水平基准量测"],
        [
            ("比较对象", "现状 GRID 与设计 GRID", "一个 CDT/TIN/GRID 与常高水平面"),
            ("积分", "每单元两个线性三角面的解析积分", "固定预算微三角形数值积分"),
            ("范围", "两幅 GRID 的共同有效全幅", "用户圈定简单多边形"),
            ("NoData", "默认整单元跳过，可由 API 开启半单元", "按微三角形中心判定有效性"),
            ("输出", "体积、面积、覆盖率、差值统计和差值 GRID", "体积、面积、覆盖率和边界 CSV"),
        ],
        [2200, 3580, 3580],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    m.table(
        ["对齐方法", "特点", "建议"],
        [
            ("双线性", "由源 GRID 的 2×2 节点连续插值；严格 NoData", "通常用于连续地形高程和设计面"),
            ("最近邻", "直接取最近源节点；不产生中间高程", "分类栅格、阶梯面或快速核查"),
        ],
        [2200, 3880, 3260],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    m.callout("符号约定", "差值 = 现状高程 − 设计高程。正值/红色计挖方，负值/蓝色的绝对体积计填方；净方 = 挖方 − 填方。若设计面统一比现状高 2 个高程单位，结果应主要为填方。", "blue")
    m.callout("坐标与精度", "程序不会隐式重投影。两幅 GRID 的 CRS WKT 不一致时直接拒绝加载设计面，应先在专业 GIS 软件中显式转换坐标系；同 CRS 错位表面则由用户明确选择最近邻或双线性对齐。土方计算采用仿射行列式得到世界 XY 面积，并解析处理三角形内部的零高差线。用于工程结算前仍应验证 CRS、单位、NoData、设计版本和行业规范。", "gold")

    m.h2("6.7 坡度、坡向与地形法向")
    for step in (
        "选择‘分析→坡度/坡向分析（单击）’；程序自动返回二维视图。",
        "在目标坡面单击一次；之后再次单击会用新位置替换上次结果。",
        "黄色圆点表示查询位置，青色虚线表示实际支撑三角形或 GRID 单元。",
        "青色箭头指向最大下降方向；箭头为固定屏幕长度，只表达方向。",
        "在右上面板核对数据源、Z、坡度角、坡向与八方位、dz/dx、dz/dy 和向上单位法向。",
        "按 Esc 或选择‘清除坡度/坡向结果’清除覆盖层。",
    ):
        m.number(step)
    m.table(
        ["结果", "定义", "注意"],
        [
            ("坡度", "相对水平面的角度，水平为 0°", "不是百分比坡度"),
            ("坡向", "最大下降方向；+Y 为北，顺时针 0～360°", "水平面无唯一坡向"),
            ("单位法向", "normalize(-dz/dx,-dz/dy,1)", "始终指向上方"),
            ("TIN/CDT", "查询三角面的解析梯度", "边/顶点采用一个有效邻面"),
            ("GRID", "2×2 支撑节点的局部双线性导数", "NoData 支撑单元拒绝分析"),
        ],
        [1900, 4260, 3200],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    m.callout("数据源规则", "每次单击按可见 CDT、TIN、GRID 优先，其后才考虑存在但隐藏的图层；选定源在该点无有效表面时本次分析失败，不用其他图层补值。可通过图层显隐明确选择叠加表面。源图层变化后结果自动清除。", "blue")
    m.callout("断裂边", "TIN/CDT 在三角面内坡度恒定。查询恰落在边或顶点时只选择一个有效邻面，不跨断裂线平均；对边两侧分别稍作偏移单击，可比较两个坡面的结果。", "gold")

    m.h2("6.8 全幅坡度、坡向与阴影专题图")
    for step in (
        "准备一个高程 GRID；若只有 TIN/CDT，直接选择专题命令，程序会自动生成最长边 401 节点的 GRID。",
        "选择‘设置专题分析与性能参数…’，依次输入 z-factor、太阳方位角、高度角、工作线程数和分块行数；一般保留线程 0、块高 0。",
        "选择‘分析→生成全幅坡度专题图’，查看绿色到紫色的 0～60° 固定色带和右上图例。",
        "选择‘生成全幅坡向专题图’，查看按北→东→南→西循环的方位色相；水平区域显示为 NoData。",
        "选择‘生成阴影地形图（当前光照参数）’，查看 0～255 灰度地形阴影。",
        "计算期间观察状态栏百分比；需要中止时按 Esc 或选择‘取消正在进行的专题计算’，旧专题结果不会被半成品替换。",
        "用‘导出当前专题图 DGRID’保存实际数值；GDAL 构建可另存为压缩 GeoTIFF。",
        "选择‘恢复显示高程 GRID’，回到原高程着色；原 GRID 始终保留。",
    ):
        m.number(step)
    m.table(
        ["专题", "数值", "屏幕表达"],
        [
            ("坡度", "相对水平面 0～90°", "固定 0～60+° 绿—黄—橙—紫"),
            ("坡向", "+Y 为北，顺时针 0～360°", "环形方位色带"),
            ("阴影", "光源点积灰度 0～255", "黑—白灰度"),
        ],
        [1800, 3600, 3960],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    m.callout("参数约定", "z-factor 默认 1.0 且必须大于 0；XY 为米、Z 为毫米时填写 0.001。太阳方位角以北为 0°顺时针并归一化到 0～360°，高度角范围 -90～90°。线程 0 自动、1 单线程、最大 64；块高 0 默认 64 行，最大 1048576。小 GRID 会自动少用线程。", "gold")
    m.callout("数据没有被截图化", "导出的 DGRID/GeoTIFF 是坡度角、坡向角或阴影值本身，不是屏幕 RGB。专题句柄独立于高程 GRID，所以三维视图、剖面、量测和局部坡面仍使用原高程。源 GRID 被替换或清除时专题结果自动释放。", "blue")
    m.callout("响应、并行与关闭", "专题任务内部按行块多线程计算，不额外复制第二份全幅缓冲；协调线程串行回报进度和检查取消。GUI 以 25 ms 周期刷新并分派窗口消息。计算期间其他菜单命令被门控；关闭窗口会先请求取消、等待全部线程退出并释放任务句柄，再完成退出。", "blue")

    m.h1("7 动态插入与删除")
    m.h2("7.1 插入点")
    for step in (
        "进入“插入模式”。",
        "在需要增加细节的位置单击。",
        "程序使用点击 XY，并按演示地形函数计算 Z。",
        "观察局部三角网即时更新及状态栏的新顶点 ID。",
    ):
        m.number(step)
    m.h2("7.2 删除点")
    for step in (
        "进入“删除模式”。",
        "在要删除的顶点附近单击。",
        "程序删除 XY 距离最近的网格顶点，而不是任意空间点。",
        "观察删除影响区和重新三角化结果。",
    ):
        m.number(step)
    m.h2("7.3 编辑颜色")
    m.table(
        ["颜色", "对象", "含义"],
        [
            ("红色", "旧三角形", "编辑前被删除或受影响的面。"),
            ("黄色", "边界边", "局部冲突/影响区域边界。"),
            ("绿色", "新三角形/新边", "编辑后生成的局部拓扑。"),
            ("洋红色", "查询三角形", "查询点所在或相交的面。"),
            ("白色", "圆形顶点标记", "查询位置的最近顶点。"),
        ],
        [1500, 2900, 4960],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    m.callout("演示数据的 Z", "GUI 插入模式的高程来自内置模拟地形函数，不能手工输入。如果要插入真实指定 Z，请由调用系统使用 DLL 的 dt_insert_point()。", "gold")

    m.h1("8 保存与打开地形数据")
    m.h2("8.1 DTMESH 文本")
    m.para("保存扩展名为 .dtmesh 或 .txt 时，程序写出 DTMESH 1 文本，包含顶点 ID、XYZ 和显式三角形索引。打开时会重建 Delaunay 并逐面校验，适合研究、交换和人工检查。")
    m.h2("8.2 DTIN 二进制")
    m.para("保存扩展名为 .dtin 时使用紧凑二进制格式。DTIN v1 保存点集和 ID，打开时重新构建三角网；共圆点可能选择另一条同样合法的对角线。")
    m.table(
        ["选择", "优点", "注意"],
        [
            ("DTMESH", "文本可读、显式面索引、加载时逐面校验", "体积较大，读写较慢。"),
            ("DTIN", "紧凑、读写快、适合程序保存", "v1 不承诺逐面拓扑完全相同。"),
        ],
        [1700, 4700, 2960],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    m.h2("8.3 DGRID 与 DCONTOUR 文本")
    m.para("通过“数据交换”菜单可独立导入或导出 DGRID 规则高程节点文本和 DCONTOUR 等高线文本。导入这些图层不会删除现有 TIN，因此适合叠加检查。GRID 支持完整六参数仿射变换；旋转或错切数据会按实际三个角点映射显示。安装目录的 sample_data/sample_grid.dgrid 和 sample_contours.dcontour 可直接用于验证。")
    m.callout("等高线反向转换", "导入 DCONTOUR 后，可执行‘等高线 → TIN’或‘等高线 → GRID’。GUI 使用原折点和 LINE elevation；输入共线或同一 XY 高程冲突时拒绝转换，原数据保持不变。", "blue")
    m.h2("8.4 GeoTIFF、COG 与 GeoPackage")
    m.para("启用 DT_WITH_GDAL 后，数据交换菜单增加五个命令：导入 GeoTIFF/COG、导出 DEFLATE GeoTIFF、导出 Cloud Optimized GeoTIFF、导入 GeoPackage 等高线和导出 GeoPackage 等高线。程序分别探测 GTiff、COG、GPKG 驱动；普通构建或缺少驱动时，相应命令置灰而不是在运行时失败。")
    m.table(
        ["对象", "导入行为", "导出设置"],
        [
            ("GeoTIFF / COG GRID", "读取第 1 波段；只替换 GRID，保留 TIN、CDT 和等高线。", "GTiff：TILED、DEFLATE、BIGTIFF=IF_SAFER；COG：DEFLATE、BIGTIFF=IF_SAFER。"),
            ("GeoPackage 等高线", "读取 elevation 与 LineStringZ；只替换等高线。", "contours 图层、elevation/closed 字段、空间索引。"),
        ],
        [2200, 3500, 3660],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    m.callout("图层安全", "导入失败时原图层保持不变。GDAL 栅格像元中心与 GRID 节点自动完成半像元换算；当前 GUI 不执行坐标重投影。", "blue")
    m.h2("8.5 DCDT 约束网")
    m.para("选择“数据交换→打开约束网 DCDT”可加载基础点、外边界、孔洞、断裂线和 CRS。域内网为紫色，外边界青色，孔洞洋红，断裂线橙色。sample_constraints.dcdt 可直接验证；“保存约束网 DCDT”执行完整文本往返。")
    m.h2("8.6 批量事务、交互绘制、移动与安全删除")
    for step in (
        "先打开 DCDT，或执行‘地形转换→从当前 TIN 创建约束网’。",
        "当前约束为空时可选择‘批量添加 12 条示例断裂线’；程序在有效范围内生成互不相交的平行线，用一次事务提交并报告耗时。",
        "选择‘约束编辑→绘制断裂线/外边界/孔洞边界’，在二维画布逐点单击。",
        "黄色草图中按 Backspace 撤销最后一点；按 Enter 完成，按 Esc 取消。",
        "断裂线至少 2 点，外边界和孔洞至少 3 点；边界由程序自动闭合。",
        "选择‘移动约束顶点（两次单击）’后，第一次在白色约束顶点 14 像素范围内单击；选中点变为黄色。",
        "第二次在目标位置单击即可提交移动；按 Esc 取消已选顶点。成功后红色表示旧面，黄色表示影响边界，绿色表示新增面/边。",
        "若移动造成未分段交叉等非法拓扑，原约束保持不变，选中状态保留，可换一个目标位置重试或按 Esc 取消。",
        "选择‘删除约束顶点（单击）’并单击白色折点；普通点直接删除并显示红/黄/绿影响。",
        "共享顶点会显示引用约束数量并请求确认；确认后只从当前约束脱离，其他约束不变。基础地形点也始终保留。",
        "删除后不足最小点数或产生非法拓扑时操作被拒绝，原约束和 generation 保持不变。",
        "选择‘选择删除约束’，在彩色约束线 14 像素范围内单击；删除外边界前先删除孔洞。",
        "约束改变后重新生成 GRID/等高线；程序会自动释放旧派生图层。",
    ):
        m.number(step)
    m.callout("高程与性能", "草图和移动目标的 Z 优先由 CDT/TIN 表面插值。v0.27 单项约束编辑仍完整重建候选 CDT，批量事务只重建一次；移动/删除顶点还请求完整差异用于影响显示。查询、插入和删除工具栏按钮仍只操作普通 TIN。", "gold")
    m.h2("8.7 往返验证")
    for step in (
        "保存当前网格。",
        "记录状态栏的顶点数和三角形数。",
        "单击“清空”。",
        "单击“打开网格”并选择刚保存的文件。",
        "比较顶点数、范围，并进行最近点和覆盖三角形查询。",
    ):
        m.number(step)

    m.h1("9 大数据量演示建议")
    m.para("DLL 对当前视口执行精确范围查询。为了避免 GDI 阻塞，二维 TIN/CDT 线框预算各约 45,000 面，三维填充与深度排序预算约 18,000 面；任意规模 GRID 先裁剪当前世界视口的源节点窗口，再生成不超过 512×512 的局部概览；等高线绘制预算约 20 万顶点。显示抽样不会删减 DLL 数据或文件导出内容。")
    m.h2("9.1 百万级点操作顺序")
    for text in (
        "先使用 10 万点确认功能和显示环境，再切换到 100 万点。",
        "生成或导入期间耐心等待，不要连续点击按钮。",
        "先框选放大到局部，再进行查询或动态编辑。",
        "避免频繁全图复位；局部范围能显示更多真实三角形细节。",
        "三维全图展示优先使用默认抽样；需要生产级连续千万点渲染时应采用 GPU 分块 LOD。",
        "超过 2000 万节点的 GRID 会按当前视口生成最多 512×512 的着色 LOD；缩放或平移结束后观察状态栏中的源窗口→预览尺寸，等高线仍从完整 GRID 计算。",
        "全幅专题计算可观察进度；确认参数或数据源不合适时及时按 Esc 取消，不会提交半成品。",
        "保存超大 DTMESH 前确认磁盘空间；需要紧凑存储时选 DTIN。",
    ):
        m.bullet(text)
    m.h2("9.2 16 GB 内存参考")
    m.para("项目实测 1,000 万随机均匀点约生成 2,000 万个有限三角形，峰值工作集约 4.55 GB，建网约 50.44 秒。实际数据分布、硬件、文件解析和同时运行的软件会影响结果，应预留额外内存。")

    m.h1("10 常见问题与恢复")
    faq = [
        ("程序启动失败", "确认 EXE、DLL、GMP 和 winpthread 四个文件位于同一 bin 目录。"),
        ("显示很稀疏", "全图抽样属于渲染策略；框选放大局部后会显示更完整细节。"),
        ("左键没有查询", "检查当前是否处于框选、插入、删除或三维视图。"),
        ("三维地形太平", "单击“高程×n”或按 + 提高垂直夸张。"),
        ("三维视点迷失", "按 Home 或单击“全图”恢复适屏相机。"),
        ("图层菜单有勾但画面没有", "勾表示显示开关；还需先生成或导入对应 GRID/等高线数据。"),
        ("编辑后 GRID 消失", "GRID/等高线是旧 TIN 的派生结果；TIN 改变后程序自动释放，需重新生成。"),
        ("GRID→TIN 提示跨 NoData", "普通 TIN 会跨空白构网；真实孔洞和硬边界请使用独立 DCDT 约束网。"),
        ("等高线→TIN 失败", "至少需要三个不共线的唯一 XY；同一位置不能属于不同高程。普通 TIN 不强制保留原等高线边。"),
        ("GDAL 菜单是灰色", "当前构建未启用 GDAL，或运行环境缺少 GTiff、COG、GPKG 对应驱动；DGRID/DCONTOUR 文本仍可使用。"),
        ("导入 GeoPackage 没有高程", "确认矢量层包含 elevation 字段和有效 LineString/LineStringZ；GUI 默认读取 elevation。"),
        ("DCDT 孔洞仍有线框", "确认显示的是紫色 CDT 域内网；底下叠加的普通 TIN 可通过图层菜单隐藏。"),
        ("无法选中约束", "删除时应靠近青、洋红或橙色约束线；移动时应靠近白色约束顶点。两者拾取范围均约 14 像素。"),
        ("外边界删除失败", "仍有孔洞依赖外边界；先删除孔洞约束，再删除外边界。"),
        ("约束提交失败", "检查点数、连续重合点和未分段交叉；失败时草图保留，可 Backspace 调整。"),
        ("约束顶点移动失败", "目标位置可能造成未分段交叉或非法边界；原约束不会改变，可再次单击其他位置，或按 Esc 取消。"),
        ("共享顶点无法删除", "这是默认保护；GUI 会询问是否只从当前约束脱离。取消时所有约束不变，确认后其他约束仍保留该点。"),
        ("约束顶点删除失败", "删除后可能少于最小点数或产生非法拓扑；原约束和基础地形点均不会被部分删除。"),
        ("批量示例被拒绝", "为避免与已有线交叉，GUI 只在当前约束为空时生成 12 条平行断裂线；可重新从 TIN 创建约束网后再试。"),
        ("剖面曲线有断点", "当前固定数据源存在 CDT 孔洞/域外、TIN 凸包外或 GRID NoData；程序不会用其他图层逐点补齐。"),
        ("剖面菜单无法导出", "先完成 A、B 两次单击并确认至少一个有效样本；源图层被编辑或替换后旧剖面会自动清除。"),
        ("坡度/坡向单击失败", "当前固定数据源在该点可能位于 CDT 孔洞/域外、TIN 凸包外或 GRID NoData 单元；调整图层显隐或在有效区域重试。"),
        ("坡向显示无唯一方向", "查询处为水平面，最大下降方向不唯一；这是有效结果，不是错误。"),
        ("断裂线附近坡度跳变", "TIN/CDT 保留分片线性表面，边上只选择一个有效邻面；在边两侧分别单击可检查两个坡面。"),
        ("专题计算耗时较长", "观察状态栏进度；按 Esc 或取消菜单可安全中止。通常使用线程 0、块高 0；CPU 受限时可指定 2～64 线程，过小块高会增加调度开销。算法仍需源 GRID 与同尺寸输出 GRID。"),
        ("阴影方向或坡度比例不正确", "使用‘设置专题分析与性能参数’核对光照方位/高度和 z-factor；XY 为米、Z 为毫米时 z-factor 应为 0.001。"),
        ("面积/土方边界被拒绝", "检查点数是否至少为 3，并去除连续重复点、自相交和退化边；按 Backspace 调整草图。"),
        ("有效覆盖率不足 100%", "固定数据源内存在 CDT 域外/孔洞、TIN 凸包外或 GRID NoData；程序不会从其他图层补齐。"),
        ("挖填方方向与预期相反", "本程序约定现状高于水平基准为挖方、低于基准为填方；净挖方等于挖方减填方。"),
        ("框选后没有变化", "选框宽或高可能小于 8 像素；重新拖出更大矩形。"),
        ("无法退出框选", "单击“查询模式”等其他模式；拖动中可按 Esc。"),
        ("导入后旧网消失", "成功导入会整体替换当前网，这是设计行为；先保存需要保留的网格。"),
        ("删除的不是点击点", "删除模式按 XY 删除最近的现有网格顶点。"),
        ("文本网格被拒绝", "DTMESH 三角形必须与其顶点重建出的 Delaunay 拓扑一致。"),
    ]
    m.table(["现象", "处理方法"], faq, [2800, 6560],
            [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT])

    m.h1("11 演示前检查清单")
    for text in (
        "确认 bin 中四个运行文件齐全。",
        "预先打开一次示例 XYZ 并保存 DTMESH，验证文件对话框和目录权限。",
        "根据电脑性能选择 10 万或 100 万模拟数据。",
        "演示查询前切换到查询模式，演示编辑前明确红/黄/绿含义。",
        "使用框选放大展示局部细节，并用“全图”恢复。",
        "依次生成 GRID 和等高线，用“图层”菜单验证 T/G/C 叠加与显隐。",
        "若有大 GRID，连续滚轮缩放、拉框和拖动平移，确认状态栏报告局部源窗口→LOD 尺寸并在松开后细化；核对导出或等高线仍来自完整数据。",
        "至少导出一次 DGRID 和 DCONTOUR，并重新导入检查范围和数量。",
        "若使用 GDAL 构建，确认五个格式菜单已启用；往返一次 GeoTIFF/COG 和 GeoPackage，并检查 TIN/CDT 未被替换。",
        "从导入的等高线反向生成一次 TIN 和 GRID，确认源等高线仍可叠加显示。",
        "打开 sample_constraints.dcdt，隐藏普通 TIN 后确认孔洞为空、三类约束颜色正确。",
        "从 TIN 创建空约束网，批量添加 12 条示例断裂线并记录状态栏事务耗时。",
        "绘制一条断裂线并按 Enter 完成，再用选择删除约束验证拾取。",
        "移动一个约束顶点，确认约束 ID 不变，并检查红色旧面、黄色边界和绿色新增面/边。",
        "删除一个普通约束顶点；若准备了共享点数据，再验证保护提示与单约束脱离。",
        "由 CDT 生成 GRID 和等高线，确认孔洞/域外没有有效 GRID 节点或跨越等高线。",
        "生成一条跨越地形的 A—B 剖面，核对曲线、统计并导出 CSV；若有孔洞，再验证断点不被其他图层补齐。",
        "圈定一块多边形，分别用两个基准高程重算并导出 CSV；若跨孔洞，再核对有效覆盖率小于 100%。",
        "保存原 GRID，复用完成的量测边界依次验证保持范围掩膜、紧凑裁剪和反向掩膜；检查 NoData、适屏、Esc 取消与重新打开恢复。",
        "打开或按高程偏移创建设计 GRID；若节点错位先执行双线性/最近邻显式对齐，再运行双表面挖填方；核对红挖蓝填、覆盖率和 CSV，并按 Esc 验证取消时旧设计仍保留。",
        "分别在 TIN/CDT 和 GRID 坡面执行单点坡度/坡向分析，核对支撑单元、下坡箭头、八方位与法向；按 Esc 清除。",
        "设置 z-factor、光照角、线程 0 和块高 0，依次生成全幅坡度、坡向和阴影专题图，检查进度与完成状态中的并行配置；再按 Esc 验证取消，导出 DGRID/专题 GeoTIFF，最后恢复高程显示。",
        "切换 3D，验证环视、滚轮缩放、WASD 漫游和垂直夸张，再返回 2D。",
        "准备原始 XYZ 备份，清空和导入会改变当前内存网格。",
    ):
        m.bullet(text)
    m.callout("推荐演示路线", "导入示例 XYZ → TIN→GRID → 连续缩放/平移并观察局部源窗口→LOD → 等高线 → 圈定多边形并作紧凑 GRID 裁剪/掩膜 → 打开错位设计 GRID → 双线性显式对齐 → 双表面挖填方、差值图与 CSV → 全幅坡度/坡向/阴影专题图与导出 → 单击坡度/坡向与法向 → 任意 A—B 剖面与 CSV → 多边形水平基准土方与 CSV → GeoTIFF/COG 与 GeoPackage 往返 → 等高线反向生成 TIN/GRID → 从 TIN 创建 CDT → 批量添加断裂线 → 绘制/移动/安全删除顶点 → CDT→GRID/等高线 → T/G/C/D 显隐 → 3D 漫游。", "blue")
    m.h2("11.1 后续 GUI 升级")
    m.para("当前已交付 2D/3D 浏览、世界视口自适应 GRID 内存概览、任意多边形节点级 GRID 裁剪/掩膜、同 CRS 错位 GRID 显式重采样、匹配现状/设计 GRID 的解析挖填方与差值图、可取消并行地形专题，以及 TIN/CDT/GRID 局部坡面、剖面、水平基准土方、格式转换、GDAL 交换及约束编辑。后续重点是坐标系显式重投影、GRID 映射文件与流式瓦片输出、持久化多级金字塔、局部 CDT 更新、网格单元解析边界裁剪、误差控制、可配置图层样式、GPU 分块渲染与贴地碰撞。")

    path = OUTPUT / "dterrain_GUI操作入门教程.docx"
    m.save(path)
    return path


if __name__ == "__main__":
    generated = [build_developer_manual(), build_gui_manual()]
    for item in generated:
        print(item)
